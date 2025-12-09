# dash/app.py
# Dashboard ECOcero – Sistema de Gestión Operativa de Calidad
# Versión: layout oscuro azul + 4 tabs analíticas + KPIs inteligentes + lectura estilo auditoría
# + Indicadores inteligentes por área/responsable con priorización de riesgo

from __future__ import annotations

import sys
from pathlib import Path
from typing import List, Dict, Any, Tuple
from io import BytesIO
from docx import Document
from docx.shared import Inches , Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import load_workbook
from openpyxl.utils import range_boundaries, get_column_letter
from datetime import datetime

import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
import numpy as np
import plotly.io as pio
import io


# ──────────────────────────────────────────────────────────────
# 1. IMPORTS Y CONFIGURACIÓN DE RUTA DEL PROYECTO
# ──────────────────────────────────────────────────────────────

from pathlib import Path
import sys
import io
import pandas as pd
from openpyxl import load_workbook, Workbook

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from db.engine import engine, _settings  # conexión SQLite + configuracion

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(root))

EXCEL_MASTER_PATH = Path(_settings["excel"]["path"]).resolve()

# ─────────────────────────────────────
# Rutas relativas de archivos clave
# ─────────────────────────────────────
EXCEL_MASTER_PATH = ROOT / "data" / "BASE DE DATOS GENERAL.xlsx"
SQLITE_PATH = ROOT / "eco8d.sqlite3"
CARPETA_INFORMES_8D = ROOT / "informes_8d"
CARPETA_INFORMES_8D.mkdir(exist_ok=True)




# ─────────────────────────────────────────────
# Pequeña protección con contraseña (demo)
# ─────────────────────────────────────────────
import os

APP_PASSWORD = os.getenv("APP_PASSWORD", "")  # lo configuraremos en Streamlit Cloud

if APP_PASSWORD:
    pwd = st.sidebar.text_input("🔒 Contraseña de acceso", type="password")
    if pwd != APP_PASSWORD:
        st.warning("Introduce la contraseña correcta para ver el panel.")
        st.stop()

# ─────────────────────────────────────────────
# ADMINISTRACIÓN DE BASE DE DATOS (sidebar)
# ─────────────────────────────────────────────
import shutil
from datetime import datetime
from pathlib import Path
from sqlalchemy import text

st.sidebar.markdown("---")
st.sidebar.markdown("### ⚙️ Administración de la base de datos")

if st.sidebar.button(
    "🧹 Vaciar tablas BDCAPA y BD8D",
    help="Crea un backup del archivo SQLite y elimina el contenido de las tablas BDCAPA y BD8D.",
    key="btn_vaciar_bd",
):
    db_path = Path("eco8d.sqlite3")

    # 1) Crear backup del archivo SQLite (si existe)
    if db_path.exists():
        backup_name = f"{db_path.stem}_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}{db_path.suffix}"
        backup_path = db_path.with_name(backup_name)
        try:
            shutil.copy2(db_path, backup_path)
            st.sidebar.info(f"Backup creado en:\n`{backup_path}`")
        except Exception as e:
            st.sidebar.warning(
                "No se pudo crear el backup del archivo SQLite.\n"
                f"Detalle: {e}"
            )

    # 2) Eliminar tablas BDCAPA y BD8D dentro del SQLite
    try:
        with engine.begin() as conn:
            conn.execute(text('DROP TABLE IF EXISTS "BDCAPA"'))
            conn.execute(text('DROP TABLE IF EXISTS "BD8D"'))

        # 3) Limpiar caché de tablas en Streamlit
        try:
            load_tables.clear()  # limpia la cache_data de load_tables()
        except Exception:
            pass

        st.sidebar.success(
            "Las tablas **BDCAPA** y **BD8D** fueron vaciadas correctamente.\n"
            "A partir de ahora irán rellenándose con los nuevos datos reales que cargues."
        )

    except Exception as e:
        st.sidebar.error(
            "No se pudo vaciar la base de datos SQLite.\n"
            f"Detalle técnico: {e}"
        )

# ──────────────────────────────────────────────────────────────
# LOGO PARA EL INFORME AUTOMÁTICO
# ──────────────────────────────────────────────────────────────
LOGO_PATH = "assets/logoECOcero.jpg"   # Ajusta si tu logo está en otro sitio

# Diccionarios globales para las figuras y sus textos en el informe
FIGS_FOR_REPORT = {}
FIGS_EXPLAIN_FOR_REPORT = {}
# ──────────────────────────────────────────────────────────────
# LOGO PARA EL INFORME AUTOMÁTICO
# ──────────────────────────────────────────────────────────────
LOGO_PATH = "assets/logoECOcero.jpg"   # Ajusta si tu logo está en otro sitio

def mostrar_fig_con_lectura(fig_key: str, titulo: str, fig, resumen: str):
    """
    Renderiza un gráfico con su 'lectura automática' debajo
    y lo registra para el informe Word.
    """

    # Mostrar gráfico con sintaxis moderna (SIN advertencias)
    st.plotly_chart(fig, width='stretch')

    # Bloque de lectura automática
    st.markdown(f"**📝 Lectura automática – {titulo}**")
    st.write(resumen)

    # Registrar para informe Word
    global FIGS_FOR_REPORT, FIGS_EXPLAIN_FOR_REPORT
    FIGS_FOR_REPORT[fig_key] = fig
    FIGS_EXPLAIN_FOR_REPORT[fig_key] = resumen
# ──────────────────────────────────────────────────────────────
# EXCEL MAESTRO — RUTAS Y HOJAS
# ──────────────────────────────────────────────────────────────
EXCEL_MASTER_PATH = Path(_settings["excel"]["path"])
EXCEL_MASTER_PATH = EXCEL_MASTER_PATH.resolve()

SHEET_8D_DB = "BDCAPAINFORME8D"   # Base de datos donde se acumulan informes
SHEET_8D_TEMPLATE = "FORMATO8D"   # Formato que usan los técnicos
SHEET_CAPA_DB = "BDCAPA"          # HOJA DONDE VIVE LA BASE DE DATOS CAPA

# Carpeta donde se guardarán copias de los informes 8D cargados
INFORMES_8D_DIR = ROOT / "informes_8d"
INFORMES_8D_DIR.mkdir(exist_ok=True)

def guardar_copia_fisica_informe(uploaded_file, numero_capa: str | None = None) -> Path:
    """
    Guarda una copia del archivo de informe 8D que subió el técnico
    en la carpeta 'informes_8d', con un nombre amigable:

        8D_CP123456_20251118_143522.xlsx

    Devuelve la ruta completa del archivo guardado.
    """
    # Limpiar nombre de CAPA para que no tenga caracteres raros
    if numero_capa:
        safe_capa = "".join(ch for ch in str(numero_capa) if ch.isalnum() or ch in "-_")
    else:
        safe_capa = "SIN_CAPA"

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")

    # Extensión original
    suffix = Path(uploaded_file.name).suffix.lower() or ".xlsx"

    filename = f"8D_{safe_capa}_{ts}{suffix}"
    dest_path = INFORMES_8D_DIR / filename

    # Volvemos al inicio del buffer por si ya se leyó antes
    uploaded_file.seek(0)
    with open(dest_path, "wb") as f:
        f.write(uploaded_file.read())

    # Dejamos el puntero otra vez al inicio por si se vuelve a usar
    uploaded_file.seek(0)
    return dest_path

# ──────────────────────────────────────────────────────────────
# EXPORTAR PLANTILLA FORMATO8D SIN PERDER DISEÑO (SOLO ESA HOJA)
# ──────────────────────────────────────────────────────────────
def preparar_plantilla_8d():
    """
    Devuelve un BytesIO con un archivo Excel que contiene
    ÚNICAMENTE la hoja FORMATO8D, copiando exactamente su diseño:

    - Colores, bordes, formatos
    - Celdas fusionadas
    - Altura de filas y ancho de columnas
    - Imágenes y otros objetos

    El archivo maestro original NO se modifica.
    """
    if not EXCEL_MASTER_PATH.exists():
        raise FileNotFoundError(
            f"No se encontró el archivo maestro en:\n{EXCEL_MASTER_PATH}"
        )

    # Cargamos el maestro
    wb_tpl = load_workbook(EXCEL_MASTER_PATH)

    if SHEET_8D_TEMPLATE not in wb_tpl.sheetnames:
        raise ValueError(
            f"La hoja '{SHEET_8D_TEMPLATE}' no existe en el archivo maestro."
        )

    # Eliminamos todas las demás hojas de ESTA copia en memoria
    for sheet_name in list(wb_tpl.sheetnames):
        if sheet_name != SHEET_8D_TEMPLATE:
            ws = wb_tpl[sheet_name]
            wb_tpl.remove(ws)

    # Guardamos esta copia (solo FORMATO8D) en memoria
    buffer_tpl = io.BytesIO()
    wb_tpl.save(buffer_tpl)
    buffer_tpl.seek(0)

    return buffer_tpl

# ──────────────────────────────────────────────────────────────
# Helper: generar un informe 8D en Word a partir de un registro
# ──────────────────────────────────────────────────────────────
def crear_doc_informe_8d(registro: pd.Series) -> BytesIO:
    """
    Crea un informe 8D en Word (docx) a partir de una fila
    de la base de datos BDCAPAINFORME8D.
    Devuelve un BytesIO listo para usar en un download_button.
    """
    doc = Document()

    # Encabezado con logo (si existe)
    try:
        if LOGO_PATH and Path(LOGO_PATH).exists():
            header = doc.sections[0].header
            p = header.paragraphs[0]
            run = p.add_run()
            run.add_picture(LOGO_PATH, width=Inches(1.5))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    except Exception:
        # Si algo falla con el logo, seguimos sin interrumpir
        pass

    # Título
    titulo = doc.add_paragraph()
    run_t = titulo.add_run("Informe 8D - ECOcero")
    run_t.bold = True
    run_t.font.size = Pt(16)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtítulo con Nº CAPA, Pedido y fecha de generación
    numero_capa = str(registro.get("Número de Incidencia (CAPA)", "") or "").strip()
    pedido = str(registro.get("Pedido", "") or "").strip()

    info_line_parts = []
    if numero_capa:
        info_line_parts.append(f"Nº Incidencia (CAPA): {numero_capa}")
    if pedido:
        info_line_parts.append(f"Pedido: {pedido}")

    info_line_parts.append(
        f"Generado el {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    )

    sub = doc.add_paragraph()
    sub_run = sub.add_run("    ·    ".join(info_line_parts))
    sub_run.font.size = Pt(10)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")  # espacio

    # Definimos el orden y etiquetas tal como en tu formulario
    campos = [
        ("Número de Incidencia (CAPA)", "Número de Incidencia (CAPA)"),
        ("Número de Reposición", "Número de Reposición"),
        ("Fecha de detección", "Fecha de detección"),
        ("Cliente", "Cliente"),
        ("Pedido", "Pedido"),
        ("Producto", "Producto"),
        ("Proceso", "Proceso"),
        ("Área", "Área"),
        ("Detectado por", "Detectado por"),
        ("Nombre1", "Nombre1"),
        ("Departamento1", "Departamento1"),
        ("Nombre2", "Nombre2"),
        ("Departamento2", "Departamento2"),
        ("Nombre3", "Nombre3"),
        ("Departamento3", "Departamento3"),
        ("¿Qué?", "¿Qué?"),
        ("¿Cuándo?", "¿Cuándo?"),
        ("¿Dónde?", "¿Dónde?"),
        ("¿Cómo?", "¿Cómo?"),
        ("¿Quién?", "¿Quién?"),
        ("¿Cuánto está impactando?", "¿Cuánto está impactando?"),
        ("Descripción Detallada de la Incidencia", "Descripción Detallada de la Incidencia"),
        ("Acción Contencion", "Acción Contencion"),
        ("Fecha de realización Contencion", "Fecha de realización Contencion "),
        ("Responsable Contencion", "Responsable Contencion"),
        ("Estatus Contencion", "Estatus Contencion"),
        ("¿Por qué? 1", "¿Por qué? 1"),
        ("¿Por qué? 2", "¿Por qué? 2"),
        ("¿Por qué? 3", "¿Por qué? 3"),
        ("¿Por qué? 4", "¿Por qué? 4"),
        ("¿Por qué? 5", "¿Por qué? 5"),
        ("Causa Raíz", "Causa Raíz"),
        ("Acción correctiva", "Acción correctiva"),
        ("Fecha correctiva", "Fecha correctiva"),
        ("Responsable correctiva", "Responsable correctiva"),
        ("Estatus Correctiva", "Estatus Correctiva"),
        ("Acción Preventiva", "Acción Preventiva"),
        ("Fecha de cierre", "Fecha de cierre"),
        ("Responsable de auditar", "Responsable de auditar"),
        ("Aprobo", "Aprobo"),
    ]

    # Creamos una tabla de 2 columnas (Campo / Valor)
    tabla = doc.add_table(rows=0, cols=2)
    tabla.style = "Table Grid"

    for etiqueta, col in campos:
        row = tabla.add_row()
        c0, c1 = row.cells

        p0 = c0.paragraphs[0]
        r0 = p0.add_run(etiqueta)
        r0.bold = True

        valor = registro.get(col, "")
        c1.text = "" if pd.isna(valor) else str(valor)

    # Pie de página simple
    doc.add_paragraph("")
    pie = doc.add_paragraph()
    pie_run = pie.add_run("Informe generado automáticamente desde el sistema ECOcero.")
    pie_run.font.size = Pt(8)
    pie.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Guardar a BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ─────────────────────────────────────────────
# AUX: leer campos del FORMATO8D (A5..A51 → B5..B51)
# ─────────────────────────────────────────────
def leer_campos_formato_8d(ws_form):
    """
    Lee el formulario vertical FORMATO8D:
    - Encabezados en columna A (filas 5 a 51)
    - Valores en columna B
    Omite filas sin interés (14, 21, 29, 34, 41, 46, 48).
    Devuelve dict: {texto_campo: valor}
    """
    filas = range(5, 52)  # 5..51 inclusive
    excluir = {14, 21, 29, 34, 41, 46, 48}
    campos = {}

    for row in filas:
        if row in excluir:
            continue
        etiqueta = ws_form[f"A{row}"].value
        valor = ws_form[f"B{row}"].value
        if etiqueta is None:
            continue
        etiqueta_str = str(etiqueta).strip()
        campos[etiqueta_str] = valor

    return campos


def _norm_text(s: str) -> str:
    """Normaliza texto para comparar: minúsculas + espacios colapsados."""
    return " ".join(str(s).strip().lower().split())

    

# ──────────────────────────────────────────────────────────────
# 2. CONFIGURACIÓN GLOBAL + CSS
# ──────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="Sistema de Gestión Operativa de Calidad",
    page_icon="📊",
    layout="wide",
)

# CSS: fondo oscuro azul eléctrico, main texto blanco, sidebar texto negro, tabs, métricas…
st.markdown(
    """
<style>
/* Fondo general de la app: azul oscuro / eléctrico */
.stApp {
    background: radial-gradient(circle at 0% 0%, #1d4ed8 0%, #020617 55%, #020617 100%);
    color: #ffffff;
    font-family: "Segoe UI", system-ui, -apple-system, BlinkMacSystemFont, sans-serif;
}

/* Contenedor principal: más padding arriba para que la tarjeta no quede pegada */
.block-container {
    padding-top: 3.2rem !important;
    padding-bottom: 3rem !important;
    max-width: 1300px;
}

/* SIDEBAR — Todo el texto en negro */
[data-testid="stSidebar"] {
    background: #f9fafb;
    color: #000000 !important;
}
[data-testid="stSidebar"] * {
    color: #000000 !important;
}

/* Títulos del sidebar */
[data-testid="stSidebar"] h1,
[data-testid="stSidebar"] h2,
[data-testid="stSidebar"] h3,
[data-testid="stSidebar"] h4,
[data-testid="stSidebar"] h5,
[data-testid="stSidebar"] h6 {
    color: #000000 !important;
    font-weight: 700;
}

/* Botones del sidebar (texto negro, fondo claro) */
[data-testid="stSidebar"] button {
    color: #000000 !important;
    border: 1px solid #47556950;
    background: #ffffff;
}

/* Inputs sidebar */
[data-testid="stSidebar"] input,
[data-testid="stSidebar"] select,
[data-testid="stSidebar"] textarea {
    color: #000000 !important;
}

/* Tablas más compactas */
.eco-table-small td, .eco-table-small th {
    font-size: 0.8rem !important;
}

/* Pestañas: texto blanco; seleccionada con borde inferior azul claro */
[data-testid="stTabs"] button {
    font-weight: 600;
    color: #e5e7eb !important;
}
[data-testid="stTabs"] button[aria-selected="true"] {
    border-bottom: 3px solid #38bdf8 !important;
    color: #ffffff !important;
}

/* Métricas */
[data-testid="stMetricValue"] {
    font-weight: 700;
    font-size: 1.1rem;
    color: #f9fafb;
}
[data-testid="stMetricDelta"] {
    color: #7dd3fc !important;
}

/* Encabezados en el main */
h1, h2, h3, h4, h5 {
    color: #ffffff;
}
p, span, label, .stMarkdown, .stCaption {
    color: #e5e7eb;
}

/* Tarjetas reutilizables */
.eco-card-kpi {
    margin-top: 1.2rem;
    padding: 1.2rem 1.5rem;
    border-radius: 18px;
    background: radial-gradient(circle at 0% 0%, #0ea5e9 0%, #0f172a 55%, #020617 100%);
    box-shadow: 0 18px 40px rgba(15, 23, 42, 0.45);
    border: 1px solid rgba(56, 189, 248, 0.45);
    
}

/* Texto pequeño dentro de tarjetas */
.eco-card p {
    font-size: 0.85rem;
}


/* Bloque de lectura / auditoría */
.eco-audit {
    background: rgba(15, 23, 42, 0.9);
    border-radius: 14px;
    padding: 0.9rem 1rem;
    border: 1px dashed rgba(148, 163, 184, 0.7);
    font-size: 0.9rem;
}

/* Ajuste del buscador global en sidebar */
.eco-search-card {
    background: linear-gradient(135deg, #e0f2fe 0%, #f9fafb 55%, #eef2ff 100%);
    border-radius: 14px;
    padding: 0.8rem 0.9rem;
    border: 1px solid rgba(148, 163, 184, 0.6);
    box-shadow: 0 6px 16px rgba(15, 23, 42, 0.25);
}
.eco-search-card label, .eco-search-card p {
    color: #020617 !important;
}

/* Tarjeta grande para el resumen de KPIs */
.eco-kpi-card {
    background:
        radial-gradient(circle at 50% 40%,
            rgba(56, 189, 248, 0.85) 0%,      /* azul eléctrico claro */
            rgba(30, 64, 175, 0.92) 45%,      /* azul eléctrico medio */
            rgba(2, 6, 23, 1) 100%            /* azul oscuro */
        ),
        linear-gradient(145deg,
            rgba(255, 255, 255, 0.14) 0%,
            rgba(255, 255, 255, 0.04) 30%,
            rgba(255, 255, 255, 0.01) 100%
        );

    border-radius: 24px;
    padding: 1.6rem 2.2rem;
    color: #f9fafb;
    margin-bottom: 1.5rem;
    width: 100%;
    box-sizing: border-box;

    /* Sombra luminosa suave alrededor */
    box-shadow:
        0 0 28px rgba(255, 255, 255, 0.32),   /* halo blanco suave */
        0 0 14px rgba(255, 255, 255, 0.22),   /* halo intermedio */
        0 18px 40px rgba(0, 0, 0, 0.45),      /* sombra normal */
        inset 0 0 12px rgba(255, 255, 255, 0.08); /* cristal interior */

    backdrop-filter: blur(4px);
}

/* Grillas internas de la tarjeta */
.eco-kpi-grid {
    display: grid;
    grid-template-columns: repeat(4, minmax(0, 1fr));
    column-gap: 1.8rem;
    row-gap: 0.6rem;
}

.eco-kpi-item {
    display: flex;
    flex-direction: column;
    justify-content: flex-start;
}

/* Etiqueta y valor */
.eco-kpi-label {
    font-size: 0.78rem;
    font-weight: 500;
    opacity: 0.9;
}

.eco-kpi-value {
    font-size: 1.25rem;
    font-weight: 700;
    margin-top: 0.15rem;
}

/* Chips de contexto (objetivo, nivel, etc.) */
.eco-kpi-chip {
    display: inline-flex;
    align-items: center;
    margin-top: 0.35rem;
    padding: 0.35rem 1.1rem;   /* ← MÁS GRANDE: alto y ancho del globo */
    border-radius: 999px;
    font-size: 0.82rem;        /* ← Texto un poquito más grande */
    font-weight: 600;
    background: rgba(15, 23, 42, 0.28);
    border: 1px solid rgba(148, 163, 184, 0.55);
}

/* Versión "verde" para nivel bajo de riesgo, etc. */
.eco-kpi-chip--green {
    background: rgba(22, 163, 74, 0.24);
    border-color: rgba(52, 211, 153, 0.85);
    color: #bbf7d0;
}
</style>
    """,
    unsafe_allow_html=True,
)

# ──────────────────────────────────────────────────────────────
# 3. FUNCIONES UTILITARIAS
# ──────────────────────────────────────────────────────────────

def _find_col_any(df: pd.DataFrame, candidates: List[str]) -> str | None:
    """Busca una columna por nombres o fragmentos (case-insensitive)."""
    if df is None or df.empty:
        return None
    cols = list(df.columns.astype(str))
    lower_map = {c.lower(): c for c in cols}

    # 1) coincidencia exacta
    for cand in candidates:
        c_low = cand.lower()
        if c_low in lower_map:
            return lower_map[c_low]

    # 2) por inclusión de fragmento
    for cand in candidates:
        c_low = cand.lower()
        for c in cols:
            if c_low in c.lower():
                return c
    return None


def find_col(df: pd.DataFrame, candidates: list[str]) -> str | None:
    return _find_col_any(df, candidates)

def _to_datetime_safe(series: pd.Series) -> pd.Series:
    """Convierte una serie a datetime sin romper si hay valores raros."""
    return pd.to_datetime(series, errors="coerce")

def _pct(n: float, d: float) -> float:
    if d <= 0:
        return 0.0
    return round(100.0 * float(n) / float(d), 1)


def to_datetime_safe(s: pd.Series) -> pd.Series:
    return pd.to_datetime(s, errors="coerce")


def compute_closure_time(df: pd.DataFrame) -> Tuple[float | None, float | None]:
    """
    Calcula tiempo promedio y mediano de cierre (en días) para las CAPA visibles.
    Usa FECHA CREACION como inicio y FECHA ALBARAN / Fecha de cierre como fin.
    """
    if df is None or df.empty:
        return None, None

    col_start = find_col(df, ["FECHA CREACION", "FECHA_CREACION", "Fecha de detección", "FECHA", "Fecha"])
    col_end = find_col(df, ["FECHA ALBARAN", "Fecha cierre", "Fecha de cierre"])

    if not col_start or not col_end:
        return None, None

    s = to_datetime_safe(df[col_start])
    e = to_datetime_safe(df[col_end])
    delta = (e - s).dt.total_seconds() / 86400.0  # días

    valid = delta.dropna()
    if valid.empty:
        return None, None

    return float(valid.mean()), float(valid.median())


def compute_sla(df: pd.DataFrame, mode: str, target: float) -> Tuple[int, int, pd.Series]:
    """
    Devuelve: (ok, ko, serie_sla)
    Si existen columnas SLA_CUMPLIDO / SLA_VENCIDO en la CAPA se usan;
    si no, se calcula por diferencias de fechas.
    """
    if df is None or df.empty:
        return 0, 0, pd.Series(dtype="float64")

    col_flag_venc = find_col(df, ["SLA_VENCIDO", "SLA VENCIDO"])
    col_flag_ok = find_col(df, ["SLA_CUMPLIDO", "SLA CUMPLIDO"])
    if col_flag_venc or col_flag_ok:
        if col_flag_ok:
            ok = int(df[col_flag_ok].fillna(False).astype(bool).sum())
            ko = len(df) - ok
        else:
            ko = int(df[col_flag_venc].fillna(False).astype(bool).sum())
            ok = len(df) - ko
        return ok, ko, pd.Series(dtype="float64")

    # Por fechas
    c_start = find_col(df, ["FECHA CREACION", "FECHA_CREACION", "Fecha de detección", "FECHA", "Fecha"])
    c_end = find_col(df, ["FECHA ALBARAN", "Fecha cierre", "Fecha de cierre"])
    if not c_start or not c_end:
        return 0, 0, pd.Series(dtype="float64")

    s = to_datetime_safe(df[c_start])
    e = to_datetime_safe(df[c_end])
    delta_sec = (e - s).dt.total_seconds()

    if mode == "días":
        sla_val = delta_sec / 86400.0
    else:
        sla_val = delta_sec / 3600.0

    ok_mask = sla_val <= target
    ok = int(ok_mask.fillna(False).sum())
    ko = len(df) - ok

    return ok, ko, pd.Series(sla_val, index=df.index)


def compute_reincidences(df: pd.DataFrame, key_a_name: str, key_b_name: str) -> Tuple[int, int, pd.DataFrame]:
    """
    Reincidencias 30/60 días (por combinación clave A + B).
    Devuelve (reinc_30, reinc_60, df_enriquecido).
    """
    if df is None or df.empty:
        return 0, 0, pd.DataFrame()

    map_keys: Dict[str, List[str]] = {
        "CLIENTE": ["CLIENTE", "Cliente", "Compañía", "Empresa"],
        "INCIDENCIA": ["INCIDENCIA", "Tipo incidencia", "Incidencia", "Motivo"],
        "PRODUCTO": ["PRODUCTO", "Producto", "Nombre del pedido", "NOMBRE DEL PEDIDO"],
        "PEDIDO": ["PEDIDO ORIGEN", "PEDIDO", "Pedido"],
        "NOMBRE DEL PEDIDO": ["NOMBRE DEL PEDIDO", "Nombre del pedido", "Producto"],
        "ÁREA": ["RESPONSABLE", "Área", "ÁREA", "AREA", "DEPARTAMENTO"],
        "RESPONSABLE": ["RESPONSABLE", "TÉCNICO", "TECNICO", "Responsable"],
    }

    c_fecha = find_col(df, ["FECHA CREACION", "FECHA_CREACION", "Fecha de detección", "FECHA", "Fecha"])
    if not c_fecha:
        return 0, 0, pd.DataFrame()

    df2 = df.copy()
    df2[c_fecha] = to_datetime_safe(df2[c_fecha])

    col_a = find_col(df2, map_keys.get(key_a_name, [key_a_name]))
    col_b = find_col(df2, map_keys.get(key_b_name, [key_b_name]))
    if not col_a or not col_b:
        return 0, 0, pd.DataFrame()

    df2 = df2.dropna(subset=[c_fecha]).sort_values([col_a, col_b, c_fecha]).reset_index(drop=True)
    df2["grp"] = df2[col_a].astype(str) + " | " + df2[col_b].astype(str)
    df2["prev_fecha"] = df2.groupby("grp")[c_fecha].shift(1)
    df2["delta_dias"] = (df2[c_fecha] - df2["prev_fecha"]).dt.days

    rec30 = int((df2["delta_dias"] <= 30).fillna(False).sum())
    rec60 = int((df2["delta_dias"] <= 60).fillna(False).sum())

    df2["reinc_30"] = (df2["delta_dias"] <= 30).fillna(False)
    df2["reinc_60"] = (df2["reinc_30"] | (df2["delta_dias"] <= 60).fillna(False))

    return rec30, rec60, df2


def make_bar_text(vals, mode: str, total: int):
    if mode == "ninguno":
        return None
    if mode == "conteo":
        return [str(int(v)) for v in vals]
    if mode == "porcentaje":
        return [f"{_pct(v, total)}%" for v in vals]
    return [f"{int(v)} ({_pct(v, total)}%)" for v in vals]


def compute_area_risk_table(
    df_view: pd.DataFrame,
    df_re: pd.DataFrame,
    col_area: str | None,
    sla_mode: str,
    sla_target: float,
) -> pd.DataFrame:
    """
    Construye una tabla de indicadores inteligentes por área / responsable:
    - Nº CAPA
    - % abiertas
    - SLA % cumplido
    - Reincidencias 30 / 60 días
    - Riesgo (0-100) y nivel de riesgo
    """
    if df_view is None or df_view.empty or not col_area or col_area not in df_view.columns:
        return pd.DataFrame()

    # Si df_re incluye la columna de área, podremos contar reincidencias por área
    has_re_area = isinstance(df_re, pd.DataFrame) and not df_re.empty and (col_area in df_re.columns)

    rows: List[Dict[str, Any]] = []
    for area_val, df_sub in df_view.groupby(col_area):
        if df_sub.empty:
            continue

        total = len(df_sub)

        # abiertas / cerradas
        col_est = find_col(df_sub, ["estatus", "status", "estado"])
        abiertas = cerradas = 0
        if col_est:
            est_vals = df_sub[col_est].astype(str).str.lower()
            cerradas = int(est_vals.str.contains("cerr", na=False).sum())
            abiertas = total - cerradas

        # SLA área
        sla_ok_a, sla_ko_a, _ = compute_sla(df_sub, sla_mode, sla_target)
        sla_tot_a = sla_ok_a + sla_ko_a
        sla_pct_a = _pct(sla_ok_a, sla_tot_a) if sla_tot_a > 0 else 0.0

        # Reincidencias por área (si se puede mapear)
        rec30_a = rec60_a = 0
        if has_re_area:
            df_re_sub = df_re[df_re[col_area] == area_val]
            if not df_re_sub.empty:
                rec30_a = int(df_re_sub["reinc_30"].sum())
                rec60_a = int(df_re_sub["reinc_60"].sum())

        # Pct abiertas
        open_pct = _pct(abiertas, total)

        # Riesgo: fórmula simple combinando % abiertas, SLA y reincidencias
        # base en 0-100
        comp_open = open_pct                       # más abiertas => más riesgo
        comp_sla = 100 - sla_pct_a                # menor SLA cumplido => más riesgo
        comp_rec = min(rec60_a * 15, 100)         # cada reincidencia suma, saturado en 100

        risk_score = round(
            0.45 * comp_open +
            0.35 * comp_sla +
            0.20 * comp_rec,
            1,
        )

        if risk_score >= 75:
            level = "ALTO"
        elif risk_score >= 50:
            level = "MEDIO"
        else:
            level = "BAJO"

        rows.append({
            "Área / Responsable": str(area_val),
            "CAPA totales": int(total),
            "% abiertas": open_pct,
            "SLA % cumplido": sla_pct_a,
            "Reincidencias 30d": int(rec30_a),
            "Reincidencias 60d": int(rec60_a),
            "Riesgo (0-100)": risk_score,
            "Nivel de riesgo": level,
        })

    if not rows:
        return pd.DataFrame()

    df_out = pd.DataFrame(rows)
    df_out = df_out.sort_values(["Riesgo (0-100)", "CAPA totales"], ascending=[False, False]).reset_index(drop=True)
    return df_out

def construir_docx_informe(cuerpo_informe: str, lista_graficos_selec):
    """
    Genera un archivo Word (DOCX) en memoria con:
    - El texto del informe (cuerpo_informe)
    - Los gráficos seleccionados (lista_graficos_selec),
      usando FIGS_FOR_REPORT y FIGS_EXPLAIN_FOR_REPORT.
    """
    # Usar las variables globales donde hemos ido guardando los gráficos
    global FIGS_FOR_REPORT, FIGS_EXPLAIN_FOR_REPORT, LOGO_PATH

    doc = Document()

    # ─────────────────────────────
    # 1) Márgenes del documento
    # ─────────────────────────────
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    # ─────────────────────────────
    # 2) Encabezado con logo
    # ─────────────────────────────
    header = doc.sections[0].header
    header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    try:
        run = header_paragraph.add_run()
        run.add_picture(LOGO_PATH, width=Inches(1.4))
    except Exception:
        # Si no encuentra la imagen, deja un texto sencillo
        header_paragraph.text = "Informe del sistema CAPA"

    # Pequeño espacio tras el encabezado
    doc.add_paragraph("")

    # ─────────────────────────────
    # 3) Título y fecha
    # ─────────────────────────────
    título = doc.add_paragraph()
    run_t = título.add_run("Informe de desempeño CAPA")
    run_t.bold = True
    run_t.font.size = Pt(16)
    título.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fecha_para = doc.add_paragraph(
        f"Fecha de generación: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    fecha_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph("")  # espacio

    # ─────────────────────────────
    # 4) Cuerpo del informe (texto)
    # ─────────────────────────────
    for bloque in cuerpo_informe.split("\n\n"):
        if bloque.strip():
            p = doc.add_paragraph(bloque)
            p.style = "Normal"

    # ─────────────────────────────
    # 5) Gráficos anexos
    # ─────────────────────────────
    if lista_graficos_selec:
        doc.add_page_break()
        doc.add_heading("Gráficos anexos", level=2)

        for idx, nombre_graf in enumerate(lista_graficos_selec, start=1):
            fig = FIGS_FOR_REPORT.get(nombre_graf)

            if fig is None:
                doc.add_paragraph(
                    f"[Aviso] No se encontró el gráfico '{nombre_graf}' "
                    "en la sesión actual del dashboard."
                )
                continue

            try:
                # Exportar a PNG en memoria (Plotly + kaleido)
                img_bytes = fig.to_image(format="png", engine="kaleido")
                img_stream = BytesIO(img_bytes)

                # Imagen centrada y ajustada al ancho de la página
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_img = p_img.add_run()
                run_img.add_picture(img_stream, width=Inches(6))  # 6" respeta márgenes

                # Caption: nombre + fecha
                cap = doc.add_paragraph(
                    f"Figura {idx}. {nombre_graf}. "
                    f"Fecha de generación: {datetime.now().strftime('%Y-%m-%d')}."
                )
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.style = "Caption"

                # Explicación corta bajo el gráfico (si la tenemos)
                explicacion = FIGS_EXPLAIN_FOR_REPORT.get(nombre_graf)
                if explicacion:
                    exp_p = doc.add_paragraph(explicacion)
                    exp_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Espacio entre figuras
                doc.add_paragraph("")

            except Exception as e:
                doc.add_paragraph(
                    f"[Error] No se pudo exportar el gráfico '{nombre_graf}': {e}"
                )

    # ─────────────────────────────
    # 6) Devolver DOCX en memoria
    # ─────────────────────────────
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generar_lectura_matriz_riesgo_prioridad(pivot_riesgo: pd.DataFrame) -> str:
    """
    Genera una lectura automática a partir de la matriz Riesgo vs Prioridad.
    pivot_riesgo: tabla con índices = Riesgo y columnas = Prioridad (conteo de CAPA)
    """
    if pivot_riesgo is None or pivot_riesgo.empty:
        return "La matriz de riesgo vs prioridad no contiene datos para analizar por el momento."

    total = pivot_riesgo.values.sum()
    if total == 0:
        return "Actualmente no hay CAPA clasificadas en la matriz de riesgo vs prioridad."

    # Encontrar la celda con más casos
    max_val = pivot_riesgo.values.max()
    # (riesgo_max, prioridad_max) será la combinación con más CAPA
    coords = (pivot_riesgo == max_val)
    riesgo_max = None
    prioridad_max = None
    for r in pivot_riesgo.index:
        for p in pivot_riesgo.columns:
            if coords.loc[r, p]:
                riesgo_max = r
                prioridad_max = p
                break
        if riesgo_max is not None:
            break

    porc_max = (max_val / total) * 100

    texto = []
    texto.append(
        f"La matriz muestra un total de **{total}** CAPA clasificadas por riesgo y prioridad."
    )

    if riesgo_max is not None and prioridad_max is not None:
        texto.append(
            f"La combinación más crítica en este momento es **riesgo '{riesgo_max}'** "
            f"con **prioridad '{prioridad_max}'**, donde se concentran **{max_val}** casos "
            f"(≈ {porc_max:.1f}% del total)."
        )

    # Resumen por riesgo
    suma_riesgo = pivot_riesgo.sum(axis=1).sort_values(ascending=False)
    top_riesgo = suma_riesgo.index[0]
    porc_top_riesgo = (suma_riesgo.iloc[0] / total) * 100

    texto.append(
        f"A nivel global de riesgo, el nivel **'{top_riesgo}'** agrupa la mayor cantidad de CAPA "
        f"(≈ {porc_top_riesgo:.1f}% del total), lo que indica dónde deberían enfocarse "
        "las primeras acciones de mejora."
    )

    return " ".join(texto)
#================================================
# funciones de analisis en tiempo real 
#================================================

def generar_lectura_prioridad(df_actions: pd.DataFrame) -> str:
    if df_actions is None or df_actions.empty or "Prioridad (automática)" not in df_actions.columns:
        return "Aún no hay datos suficientes para analizar la distribución por prioridad."

    series = df_actions["Prioridad (automática)"].value_counts()
    total = int(series.sum())
    if total == 0:
        return "Actualmente no hay CAPA con prioridad asignada."

    partes = [f"Se han analizado **{total}** CAPA con prioridad automática. "]

    for nivel in ["Alta", "Media", "Baja"]:
        if nivel in series.index:
            n = int(series[nivel])
            partes.append(f"{nivel}: {n} casos ({n/total*100:.1f}%). ")

    # Detectar si hay sobrecarga en prioridad alta
    if "Alta" in series.index and series["Alta"] / total > 0.4:
        partes.append(
            "Existe una concentración importante en prioridad **Alta**, lo que sugiere "
            "revisar capacidad de respuesta y recursos disponibles."
        )
    else:
        partes.append(
            "La distribución entre prioridades parece relativamente equilibrada, "
            "sin una sobrecarga extrema en prioridad Alta."
        )

    return "".join(partes)


def generar_lectura_riesgo(df_actions: pd.DataFrame) -> str:
    if df_actions is None or df_actions.empty or "Riesgo (automático)" not in df_actions.columns:
        return "Aún no hay datos suficientes para analizar la distribución por nivel de riesgo."

    series = df_actions["Riesgo (automático)"].value_counts()
    total = int(series.sum())
    if total == 0:
        return "Actualmente no hay CAPA con riesgo calculado."

    partes = [f"En cuanto al nivel de riesgo, se han clasificado **{total}** CAPA. "]

    for nivel in ["Muy alto", "Alto", "Medio", "Bajo"]:
        if nivel in series.index:
            n = int(series[nivel])
            partes.append(f"{nivel}: {n} casos ({n/total*100:.1f}%). ")

    # Comentario global
    porc_critico = 0.0
    for nivel in ["Muy alto", "Alto"]:
        if nivel in series.index:
            porc_critico += series[nivel] / total * 100

    if porc_critico >= 30:
        partes.append(
            f"Alrededor de {porc_critico:.1f}% de las CAPA están en niveles de riesgo "
            "Altos o Muy altos, lo que indica un nivel de exposición relevante "
            "que debería tratarse de forma prioritaria en los comités de seguimiento."
        )
    else:
        partes.append(
            f"Solo alrededor de {porc_critico:.1f}% de los casos están en riesgo Alto/Muy alto; "
            "la mayoría se concentra en niveles Medio o Bajo."
        )

    return "".join(partes)


# ──────────────────────────────────────────────────────────────
# 4. CARGA DE TABLAS DESDE SQLITE (CACHE)
# ──────────────────────────────────────────────────────────────

@st.cache_data(show_spinner=False, ttl=60)
def load_tables() -> Dict[str, pd.DataFrame]:
    """Carga BDCAPA y BD8D desde SQLite."""
    out: Dict[str, pd.DataFrame] = {}
    with engine.connect() as conn:
        try:
            df_capa = pd.read_sql('SELECT * FROM "BDCAPA"', conn)
            out["bdcapa"] = df_capa
        except Exception:
            out["bdcapa"] = pd.DataFrame()

        try:
            df_8d = pd.read_sql('SELECT * FROM "BD8D"', conn)
            out["bd8d"] = df_8d
        except Exception:
            out["bd8d"] = pd.DataFrame()
    return out


# ──────────────────────────────────────────────────────────────
# 5. SIDEBAR – CONTROLES, PLANTILLAS, FILTROS, BUSCADOR, 8D
# ──────────────────────────────────────────────────────────────

# ─────────────────────────────────────────────
# 5.1 Plantilla informe 8D (para técnicos)
# ─────────────────────────────────────────────
st.sidebar.markdown("### 📑 Plantilla informe 8D (para técnicos)")

try:
    plantilla_bytes = preparar_plantilla_8d()
    st.sidebar.download_button(
        label="⬇️ Descargar plantilla FORMATO8D",
        data=plantilla_bytes,
        file_name="FORMATO8D_plantilla.xlsx",
        mime=(
            "application/vnd.openxmlformats-officedocument."
            "spreadsheetml.sheet"
        ),
        key="dl_formato_8d_sidebar",
        help=(
            "Descarga una copia idéntica al FORMATO8D original "
            "(colores, fusionados, bordes y tamaños intactos)."
        ),
    )
except FileNotFoundError:
    st.sidebar.warning(
        "No se encontró **'BASE DE DATOS GENERAL.xlsx'** en la carpeta del dashboard. "
        "Colócalo en el mismo directorio donde está `data`."
    )
except ValueError as e:
    st.sidebar.warning(str(e))
except Exception as e:
    st.sidebar.warning(
        "No se pudo preparar la plantilla 8D. "
        "Revise que el archivo **'BASE DE DATOS GENERAL.xlsx'** existe "
        f"y contiene la hoja **'{SHEET_8D_TEMPLATE}'**.\n\n"
        f"Detalle técnico: {e}"
    )

# ─────────────────────────────────────────────
# 5.2 Controles generales
# ─────────────────────────────────────────────

st.sidebar.markdown("### ⚙️ Controles")

if st.sidebar.button("🔄 Actualizar datos", key="btn_refresh_data"):
    # Limpiamos la caché de tablas y forzamos un rerun completo
    try:
        load_tables.clear()
    except Exception:
        pass
    st.rerun()

# Cargar datos desde SQLite
try:
    tables = load_tables()
    df_capa = tables.get("bdcapa", pd.DataFrame())
    df_8d = tables.get("bd8d", pd.DataFrame())
except Exception as e:
    st.sidebar.error(f"Error cargando tablas desde SQLite: {e}")
    st.stop()

# ─────────────────────────────────────────────
# 5.3 Estado de datos
# ─────────────────────────────────────────────
st.sidebar.markdown("---")
st.sidebar.markdown("### 🗂️ Estado de datos")

n_capa = int(len(df_capa)) if df_capa is not None else 0
n_8d = int(len(df_8d)) if df_8d is not None else 0
st.sidebar.write(f"**BDCAPA** : {n_capa} registros")
st.sidebar.write(f"**BD8D** : {n_8d} registros")

# ─────────────────────────────────────────────
# 5.4 Filtros básicos sobre BDCAPA
# ─────────────────────────────────────────────
# Detectar columnas clave en BDCAPA
col_capa = _find_col_any(df_capa, ["nº capa", "no capa", "numero capa", "n cap", "capa"])
col_estatus = _find_col_any(df_capa, ["estatus", "status", "estado"])
col_cliente = _find_col_any(df_capa, ["cliente"])
col_comercial = _find_col_any(df_capa, ["comercial"])
col_tecnico = _find_col_any(df_capa, ["tecnico", "técnico"])
col_responsable = _find_col_any(df_capa, ["responsable", "RESPONSABLE"])
col_acc_cont = _find_col_any(
    df_capa,
    [
        "accion contención",
        "accion contencion",
        "acc. cont",
        "acc cont",
        "ACCION CONTENCIÓN",
        "ACCION CONTENCION",
    ],
)
col_acc_corr = _find_col_any(
    df_capa,
    [
        "accion corrección",
        "accion correccion",
        "acc. corr",
        "acc corr",
        "ACCION CORRECCIÓN",
        "ACCION CORRECCION",
    ],
)
col_cost_rep = _find_col_any(
    df_capa,
    ["€ reposición", "€ reposicion", "costo reposicion", "e reposicion"],
)
col_cost_dev = _find_col_any(
    df_capa,
    ["€ devolucion", "€ devolución", "costo devolucion", "e devolucion"],
)

st.sidebar.markdown("---")
st.sidebar.markdown("### 🔍 Filtros CAPA")

df_view = df_capa.copy()

if col_cliente and not df_view.empty:
    clientes = sorted(df_view[col_cliente].dropna().astype(str).unique())
    sel_cli = st.sidebar.multiselect("Cliente", clientes)
    if sel_cli:
        df_view = df_view[df_view[col_cliente].astype(str).isin(sel_cli)]

if col_tecnico and not df_view.empty:
    tecnicos = sorted(df_view[col_tecnico].dropna().astype(str).unique())
    sel_tec = st.sidebar.multiselect("Técnico", tecnicos)
    if sel_tec:
        df_view = df_view[df_view[col_tecnico].astype(str).isin(sel_tec)]

if col_estatus and not df_view.empty:
    ests = sorted(df_view[col_estatus].dropna().astype(str).unique())
    sel_est = st.sidebar.multiselect("Estatus", ests)
    if sel_est:
        df_view = df_view[df_view[col_estatus].astype(str).isin(sel_est)]

st.sidebar.markdown("---")
st.sidebar.caption("Filtros aplicados sólo sobre **BDCAPA** (vista operativa).")

# ─────────────────────────────────────────────
# 5.5 Opciones avanzadas (KPI / SLA / Reincidencias)
# ─────────────────────────────────────────────
st.sidebar.markdown("### ⚙️ Opciones KPI / SLA / Reincidencias")
with st.sidebar.expander("Mostrar opciones de visualización", expanded=False):
    label_mode = st.selectbox(
        "Etiquetas en barras",
        ["conteo", "porcentaje", "ambos", "ninguno"],
        index=2,
    )
    sla_mode = st.selectbox("Unidad SLA", ["días", "horas"], index=0)
    sla_target = st.number_input(
        "Objetivo SLA (>=0)",
        min_value=0.0,
        value=7.0,
        step=1.0,
    )

    re_key_a = st.selectbox(
        "Clave A",
        ["CLIENTE", "INCIDENCIA", "PRODUCTO", "PEDIDO", "NOMBRE DEL PEDIDO"],
        index=0,
    )
    re_key_b = st.selectbox(
        "Clave B",
        ["INCIDENCIA", "PRODUCTO", "CLIENTE", "ÁREA", "RESPONSABLE"],
        index=1,
    )

# ─────────────────────────────────────────────
# 5.6 Buscador global CAPA + 8D
# ─────────────────────────────────────────────
st.sidebar.markdown("---")
st.sidebar.markdown("### 🔎 Buscador global")

with st.sidebar.container():
    st.markdown('<div class="eco-search-card">', unsafe_allow_html=True)
    search_text = st.text_input(
        "Buscar texto en CAPA (vista filtrada) y 8D",
        value="",
        placeholder="Ej: cliente, incidencia, responsable, nº CAPA…",
    )
    st.caption(
        "Filtra ambas tablas por coincidencias de texto en cualquier columna "
        "(no distingue mayúsculas/minúsculas)."
    )
    st.markdown("</div>", unsafe_allow_html=True)

df_view_search = df_view.copy()
df_8d_view = df_8d.copy()

if search_text and not df_view_search.empty:
    s = str(search_text).strip().lower()
    mask = pd.Series(False, index=df_view_search.index)
    for c in df_view_search.columns:
        mask = mask | df_view_search[c].astype(str).str.lower().str.contains(s, na=False)
    df_view_search = df_view_search[mask]

if search_text and not df_8d_view.empty:
    s = str(search_text).strip().lower()
    mask8 = pd.Series(False, index=df_8d_view.index)
    for c in df_8d_view.columns:
        mask8 = mask8 | df_8d_view[c].astype(str).str.lower().str.contains(s, na=False)
    df_8d_view = df_8d_view[mask8]

# ─────────────────────────────────────────────
# 5.7 Carga automática de informes 8D → BDCAPAINFORME8D + cruce con BDCAPA
# ─────────────────────────────────────────────
st.sidebar.markdown("---")
st.sidebar.markdown("### 📥 Cargar informe de inspección (8D)")

uploaded_8d = st.sidebar.file_uploader(
    "Selecciona el informe de inspección (8D)",
    type=["xlsx", "xls"],
    key="upl_informe_8d",
    help="Sube el informe que genera el técnico, basado en la plantilla FORMATO8D.",
)

procesar_8d = st.sidebar.button(
    "➕ Incorporar informe a BDCAPAINFORME8D",
    key="btn_merge_informe_8d",
    help="Lee los datos de la hoja FORMATO8D y agrega UNA sola fila a la base de datos.",
)


def _last_data_row_in_table(ws, min_row, min_col, max_col) -> int:
    """
    Devuelve la última fila (>= min_row) que tiene al menos
    una celda con datos entre min_col y max_col.
    Si no hay datos, devuelve min_row - 1.
    """
    last = min_row - 1
    max_sheet_row = ws.max_row
    for r in range(min_row, max_sheet_row + 1):
        row_has_data = False
        for c in range(min_col, max_col + 1):
            v = ws.cell(row=r, column=c).value
            if v not in (None, ""):
                row_has_data = True
                break
        if row_has_data:
            last = r
    return last


if procesar_8d:
    if uploaded_8d is None:
        st.sidebar.warning("Primero selecciona un archivo de informe para procesar.")
    else:
        st.sidebar.info(f"Procesando archivo de informe: **{uploaded_8d.name}**")

        # 1) Abrir el informe cargado y leer directamente la hoja FORMATO8D
        try:
            wb_informe = load_workbook(uploaded_8d, data_only=True)
        except Exception as e:
            st.sidebar.error(f"No se pudo leer el archivo del informe. Detalle: {e}")
            wb_informe = None

        if wb_informe is not None:
            # Intentar usar la hoja FORMATO8D; si no está, usar la activa
            if SHEET_8D_TEMPLATE in wb_informe.sheetnames:
                ws_form = wb_informe[SHEET_8D_TEMPLATE]
            else:
                ws_form = wb_informe.active
                st.sidebar.warning(
                    f"No se encontró la hoja '{SHEET_8D_TEMPLATE}' en el archivo subido. "
                    "Se usó la hoja activa del libro."
                )

            # 2) Construir UN solo registro leyendo las celdas B5, B6, B7… B51
            registro = {
                "Número de Incidencia (CAPA)": ws_form["B5"].value,
                "Número de Reposición": ws_form["B6"].value,
                "Fecha de detección": ws_form["B7"].value,
                "Cliente": ws_form["B8"].value,
                "Pedido": ws_form["B9"].value,
                "Producto": ws_form["B10"].value,
                "Proceso": ws_form["B11"].value,
                "Área": ws_form["B12"].value,
                "Detectado por": ws_form["B13"].value,
                "Nombre1": ws_form["B15"].value,
                "Departamento1": ws_form["B16"].value,
                "Nombre2": ws_form["B17"].value,
                "Departamento2": ws_form["B18"].value,
                "Nombre3": ws_form["B19"].value,
                "Departamento3": ws_form["B20"].value,
                "¿Qué?": ws_form["B22"].value,
                "¿Cuándo?": ws_form["B23"].value,
                "¿Dónde?": ws_form["B24"].value,
                "¿Cómo?": ws_form["B25"].value,
                "¿Quién?": ws_form["B26"].value,
                "¿Cuánto está impactando?": ws_form["B27"].value,
                "Descripción Detallada de la Incidencia": ws_form["B28"].value,
                "Acción Contencion": ws_form["B30"].value,
                "Fecha de realización Contencion ": ws_form["B31"].value,
                "Responsable Contencion": ws_form["B32"].value,
                "Estatus Contencion": ws_form["B33"].value,
                "¿Por qué? 1": ws_form["B35"].value,
                "¿Por qué? 2": ws_form["B36"].value,
                "¿Por qué? 3": ws_form["B37"].value,
                "¿Por qué? 4": ws_form["B38"].value,
                "¿Por qué? 5": ws_form["B39"].value,
                "Causa Raíz": ws_form["B40"].value,
                "Acción correctiva": ws_form["B42"].value,
                "Fecha correctiva": ws_form["B43"].value,
                "Responsable correctiva": ws_form["B44"].value,
                "Estatus Correctiva": ws_form["B45"].value,
                "Acción Preventiva": ws_form["B47"].value,
                "Fecha de cierre": ws_form["B49"].value,
                "Responsable de auditar": ws_form["B50"].value,
                "Aprobo": ws_form["B51"].value,
            }

            # 2B) Guardar copia física del archivo 8D cargado
            numero_capa_str = str(registro.get("Número de Incidencia (CAPA)", "") or "").strip()
            try:
                ruta_copia = guardar_copia_fisica_informe(
                    uploaded_8d,
                    numero_capa_str if numero_capa_str else None,
                )
                st.sidebar.info(
                    "Se guardó una copia del informe 8D en:\n\n"
                    f"`{ruta_copia}`"
                )
            except Exception as e:
                ruta_copia = None
                st.sidebar.warning(
                    "El registro 8D se seguirá incorporando a la base de datos, "
                    f"pero no se pudo guardar la copia física del archivo.\n\nDetalle: {e}"
                )

            # 3) Abrir el Excel maestro
            if not EXCEL_MASTER_PATH.exists():
                st.sidebar.error(
                    f"No se encontró el archivo maestro: {EXCEL_MASTER_PATH}\n"
                    "Verifica el nombre o la ruta del Excel general."
                )
            else:
                try:
                    wb_master = load_workbook(EXCEL_MASTER_PATH)
                except Exception as e:
                    st.sidebar.error(f"No se pudo abrir el Excel maestro. Detalle: {e}")
                    wb_master = None

                if wb_master is not None:
                    # ─────────────────────────────────────
                    # 3A) Actualizar hoja BDCAPAINFORME8D
                    # ─────────────────────────────────────
                    if SHEET_8D_DB not in wb_master.sheetnames:
                        st.sidebar.error(
                            f"No se encontró la hoja de base de datos **'{SHEET_8D_DB}'** "
                            "dentro del Excel maestro."
                        )
                    else:
                        ws_db = wb_master[SHEET_8D_DB]

                        # Encabezados de la fila 2 (tabla de informes 8D)
                        headers_raw = [c.value or "" for c in ws_db[2]]

                        # Normalizar diccionario de registro (minúsculas)
                        registro_norm_8d = {}
                        for k, v in registro.items():
                            if k is None:
                                continue
                            k_norm = str(k).strip().lower()
                            registro_norm_8d[k_norm] = v

                        # Construir la fila nueva respetando el orden de columnas de la tabla
                        fila_nueva_8d = []
                        for h in headers_raw:
                            h_text = str(h) if h is not None else ""
                            h_norm = h_text.strip().lower()
                            fila_nueva_8d.append(registro_norm_8d.get(h_norm, None))

                        # Insertar la fila dentro de la tabla SIN dejar filas vacías
                        if ws_db.tables:
                            tbl = list(ws_db.tables.values())[0]
                            min_col, min_row, max_col, max_row = range_boundaries(tbl.ref)

                            # la primera fila de datos es min_row + 1 (si min_row es encabezado)
                            first_data_row = min_row + 1
                            last_data_row = _last_data_row_in_table(
                                ws_db, first_data_row, min_col, max_col
                            )
                            new_row = last_data_row + 1 if last_data_row >= first_data_row else first_data_row

                            # Escribir la fila NUEVA en new_row
                            for idx, value in enumerate(fila_nueva_8d, start=min_col):
                                ws_db.cell(row=new_row, column=idx, value=value)

                            # Actualizar rango de la tabla
                            new_ref = (
                                f"{get_column_letter(min_col)}{min_row}:"
                                f"{get_column_letter(max_col)}{new_row}"
                            )
                            tbl.ref = new_ref
                        else:
                            ws_db.append(fila_nueva_8d)

                    # ─────────────────────────────────────
                    # 3B) Cruce con BDCAPA por Nº de Reposición
                    # ─────────────────────────────────────
                    from unicodedata import normalize as _uni_normalize
                    import re as _re

                    def _norm(s: str) -> str:
                        if s is None:
                            return ""
                        s = str(s).strip().lower()
                        # quitar acentos
                        s = "".join(
                            c for c in _uni_normalize("NFD", s)
                            if not (ord(c) >= 0x300 and ord(c) <= 0x036F)
                        )
                        # reemplazar caracteres raros por espacio
                        s = _re.sub(r"[^a-z0-9 ]+", " ", s)
                        s = _re.sub(r"\s+", " ", s).strip()
                        return s

                    # Normalizar diccionario de registro (FORMATO 8D)
                    registro_norm = {}
                    for k, v in registro.items():
                        if k is None:
                            continue
                        registro_norm[_norm(k)] = v

                    # Valor de Nº de Reposición del informe
                    repo_val = registro.get("Número de Reposición")
                    repo_val_str = str(repo_val).strip() if repo_val is not None else ""

                    # Si no hay Nº de Reposición, NO hacemos cruce con BDCAPA (para evitar filas vacías)
                    if not repo_val_str:
                        st.sidebar.info(
                            "El informe no tiene Nº de Reposición. Se guardó en BDCAPAINFORME8D, "
                            "pero no se cruzó con BDCAPA."
                        )
                    else:
                        if SHEET_CAPA_DB in wb_master.sheetnames:
                            ws_capa = wb_master[SHEET_CAPA_DB]

                            # Encabezados de la fila 2 de BDCAPA
                            headers_capa_raw = [c.value or "" for c in ws_capa[2]]
                            header_norm_to_idx = {}
                            for idx, h in enumerate(headers_capa_raw, start=1):
                                h_text = str(h) if h is not None else ""
                                header_norm_to_idx[_norm(h_text)] = idx

                            # Mapa explícito: BDCAPA → FORMATO 8D (normalizados)
                            campo_map = {
                                "n reposicion": "numero de reposicion",
                                "numero de reposicion": "numero de reposicion",
                                "fecha creacion": "fecha de deteccion",
                                "cliente": "cliente",
                                "pedido origen": "pedido",
                                "nombre del pedido": "pedido",
                                "tecnico": "detectado por",
                                "incidencia": "descripcion detallada de la incidencia",
                                "responsable": "responsable correctiva",
                                "accion contencion": "accion contencion",
                                "accion correccion": "accion correctiva",
                                "descripcion resumida causa raiz extendida en informe 8d": "causa raiz",
                                "enlace informe 8d": "__enlace_8d__",  # especial: usar ruta_copia
                            }

                            # Detectar columna Nº REPOSICIÓN en BDCAPA
                            col_idx_repo = None
                            for cand in ["n reposicion", "numero de reposicion"]:
                                cand_norm = _norm(cand)
                                if cand_norm in header_norm_to_idx:
                                    col_idx_repo = header_norm_to_idx[cand_norm]
                                    break

                            # Detectar columna Nº CAPA en BDCAPA
                            col_idx_capa = None
                            for cand in [
                                "n capa",
                                "numero de capa",
                                "no capa",
                                "capa",
                                "numero de incidencia capa",
                                "numero de incidencia (capa)",
                            ]:
                                cand_norm = _norm(cand)
                                if cand_norm in header_norm_to_idx:
                                    col_idx_capa = header_norm_to_idx[cand_norm]
                                    break

                            row_match = None
                            max_row_capa = ws_capa.max_row

                            # 1) Buscar fila existente por Nº de Reposición
                            if repo_val_str and col_idx_repo is not None:
                                for row_idx in range(3, max_row_capa + 1):
                                    cell_val = ws_capa.cell(row=row_idx, column=col_idx_repo).value
                                    if str(cell_val).strip() == repo_val_str:
                                        row_match = row_idx
                                        break

                            def _valor_para_campo_capa(capa_header_text: str):
                                """
                                Dado el nombre de la columna en BDCAPA,
                                devuelve el valor que debemos escribir (a partir del 8D + reglas).
                                """
                                capa_norm = _norm(capa_header_text)

                                # Estatus por defecto para nuevas CAPA
                                if capa_norm == "estatus":
                                    return "ABIERTA"

                                # Enlace informe 8D (ruta de la copia física)
                                if capa_norm == "enlace informe 8d":
                                    return str(ruta_copia) if ruta_copia is not None else ""

                                origen_norm = campo_map.get(capa_norm, None)

                                if origen_norm == "__enlace_8d__":
                                    return str(ruta_copia) if ruta_copia is not None else ""

                                if origen_norm is None:
                                    # Si no hay mapeo explícito, intentamos usar el mismo nombre
                                    origen_norm = capa_norm

                                return registro_norm.get(origen_norm, None)

                            # 2) Si no existe fila con ese Nº de Reposición → crear nueva CAPA
                            if row_match is None:
                                # Generar nuevo Nº CAPA si hay columna CAPA
                                nuevo_capa = None
                                if col_idx_capa is not None:
                                    max_num = 0
                                    for row_idx in range(3, max_row_capa + 1):
                                        val = ws_capa.cell(row=row_idx, column=col_idx_capa).value
                                        if val is None:
                                            continue
                                        try:
                                            s_val = str(val).strip().upper()  # CP455, CP-455, CP 455, etc.

                                        # Quitar prefijo CP / CP- / CP (con espacio)
                                            if s_val.startswith("CP-"):
                                                s_val = s_val[3:]
                                            elif s_val.startswith("CP "):
                                                s_val = s_val[3:]
                                            elif s_val.startswith("CP"):
                                                s_val = s_val[2:]

                                            # Quitar posibles espacios intermedios (por si viene "455 " o " 455")
                                            s_val = s_val.replace(" ", "")

                                            n = int(s_val)
                                            if n > max_num:
                                                max_num = n
                                        except (TypeError, ValueError):
                                            # Si no se puede convertir, se ignora ese valor
                                            continue

                                    nuevo_capa = max_num + 1

                                # Construir fila nueva para BDCAPA
                                fila_capa = []
                                for h in headers_capa_raw:
                                    h_text = str(h) if h is not None else ""
                                    h_norm = _norm(h_text)
                                    idx_actual = header_norm_to_idx.get(h_norm)

                                    if (
                                        col_idx_capa is not None
                                        and idx_actual == col_idx_capa
                                        and nuevo_capa is not None
                                    ):
                                        # Nº CAPA con prefijo CP
                                        valor = f"CP{nuevo_capa}"
                                    else:
                                        valor = _valor_para_campo_capa(h_text)

                                    fila_capa.append(valor)

                                # Insertar dentro de la tabla de BDCAPA si existe, sin filas vacías intermedias
                                if ws_capa.tables:
                                    tbl_capa = list(ws_capa.tables.values())[0]
                                    c_min_col, c_min_row, c_max_col, c_max_row = range_boundaries(tbl_capa.ref)

                                    first_data_row_capa = c_min_row + 1
                                    last_data_row_capa = _last_data_row_in_table(
                                        ws_capa, first_data_row_capa, c_min_col, c_max_col
                                    )
                                    new_row_capa = (
                                        last_data_row_capa + 1
                                        if last_data_row_capa >= first_data_row_capa
                                        else first_data_row_capa
                                    )

                                    for idx, value in enumerate(fila_capa, start=c_min_col):
                                        ws_capa.cell(row=new_row_capa, column=idx, value=value)

                                    new_ref_capa = (
                                        f"{get_column_letter(c_min_col)}{c_min_row}:"
                                        f"{get_column_letter(c_max_col)}{new_row_capa}"
                                    )
                                    tbl_capa.ref = new_ref_capa
                                else:
                                    ws_capa.append(fila_capa)

                            else:
                                # 3) Si ya existe la fila con ese Nº de Reposición → actualizar campos (sin tocar Nº CAPA)
                                for h_idx, h in enumerate(headers_capa_raw, start=1):
                                    h_text = str(h) if h is not None else ""
                                    h_norm = _norm(h_text)

                                    # No sobreescribimos Nº CAPA
                                    if col_idx_capa is not None and h_idx == col_idx_capa:
                                        continue

                                    nuevo_valor = _valor_para_campo_capa(h_text)

                                    # Solo escribir si trae algo (no None / vacío)
                                    if nuevo_valor not in (None, ""):
                                        ws_capa.cell(row=row_match, column=h_idx, value=nuevo_valor)

                        else:
                            st.sidebar.warning(
                                f"No se encontró la hoja **'{SHEET_CAPA_DB}'** en el Excel maestro. "
                                "No se pudo actualizar la base de datos CAPA."
                            )

                    # 4) Guardar Excel maestro y sincronizar con SQLite
                    try:
                        wb_master.save(EXCEL_MASTER_PATH)
                    except PermissionError:
                        st.sidebar.error(
                            "No se pudo guardar el Excel maestro (permiso denegado).\n"
                            "Cierra 'BASE DE DATOS GENERAL.xlsx' si lo tienes abierto "
                            "y vuelve a intentarlo."
                        )
                    else:
                        # 4A) Actualizar BD8D en SQLite
                        try:
                            df_raw_8d = pd.read_excel(
                                EXCEL_MASTER_PATH,
                                sheet_name=SHEET_8D_DB,
                                header=None,
                            )
                            header_row_8d = None
                            for i in range(len(df_raw_8d)):
                                if df_raw_8d.iloc[i].count() >= 5:
                                    header_row_8d = i
                                    break

                            if header_row_8d is not None:
                                df_8d_db = pd.read_excel(
                                    EXCEL_MASTER_PATH,
                                    sheet_name=SHEET_8D_DB,
                                    header=header_row_8d,
                                )
                                with engine.begin() as conn:
                                    df_8d_db.to_sql("BD8D", conn, if_exists="replace", index=False)
                            else:
                                st.sidebar.warning(
                                    "El Excel se actualizó, pero no se detectaron correctamente "
                                    "los encabezados en BDCAPAINFORME8D. No se actualizó BD8D en SQLite."
                                )
                        except Exception as e:
                            st.sidebar.warning(
                                "El Excel se actualizó correctamente, "
                                "pero no se pudo actualizar la tabla BD8D en SQLite.\n"
                                f"Detalle técnico: {e}"
                            )

                        # 4B) Actualizar BDCAPA en SQLite (si existe la hoja)
                        try:
                            if SHEET_CAPA_DB in wb_master.sheetnames:
                                df_raw_capa = pd.read_excel(
                                    EXCEL_MASTER_PATH,
                                    sheet_name=SHEET_CAPA_DB,
                                    header=None,
                                )
                                header_row_capa = None
                                for i in range(len(df_raw_capa)):
                                    if df_raw_capa.iloc[i].count() >= 5:
                                        header_row_capa = i
                                        break

                                if header_row_capa is not None:
                                    df_capa_db = pd.read_excel(
                                        EXCEL_MASTER_PATH,
                                        sheet_name=SHEET_CAPA_DB,
                                        header=header_row_capa,
                                    )
                                    with engine.begin() as conn:
                                        df_capa_db.to_sql("BDCAPA", conn, if_exists="replace", index=False)
                                else:
                                    st.sidebar.warning(
                                        "El Excel se actualizó, pero no se detectaron correctamente "
                                        "los encabezados en BDCAPA. No se actualizó BDCAPA en SQLite."
                                    )
                        except Exception as e:
                            st.sidebar.warning(
                                "El Excel se actualizó correctamente, "
                                "pero no se pudo actualizar la tabla BDCAPA en SQLite.\n"
                                f"Detalle técnico: {e}"
                            )

                        # Limpiar caché de tablas para que el sidebar vea los datos nuevos
                        try:
                            load_tables.clear()
                        except Exception:
                            pass

                        st.sidebar.success(
                            f"Informe incorporado correctamente en **'{SHEET_8D_DB}'** "
                            f"y cruzado con **'{SHEET_CAPA_DB}'** (si tenía Nº de Reposición).\n"
                            "Se ha actualizado la tabla 8D y la base CAPA."
                        )
                        # Limpiar caché de tablas para que el sidebar vea los datos nuevos
                        try:
                            load_tables.clear()
                        except Exception:
                            pass

                        st.rerun()
                        
# ─────────────────────────────────────────────
# 5.8 Generar informe 8D (Word) desde BDCAPAINFORME8D por Nº de Pedido
# ─────────────────────────────────────────────
st.sidebar.markdown("---")
st.sidebar.markdown("### 📄 Generar informe 8D desde la base de datos")

try:
    # Leer sin encabezado
    df_raw = pd.read_excel(
        EXCEL_MASTER_PATH,
        sheet_name=SHEET_8D_DB,
        header=None
    )

    # Detectar la fila que contiene los encabezados reales
    header_row = None
    for i in range(len(df_raw)):
        if df_raw.iloc[i].count() >= 5:   # fila con al menos 5 celdas no vacías
            header_row = i
            break

    if header_row is None:
        st.sidebar.warning(
            "No se pueden detectar los encabezados reales en BDCAPAINFORME8D. "
            "No es posible generar informes."
        )
        df_8d_db = pd.DataFrame()
    else:
        # Leer de nuevo usando la fila detectada como encabezado
        df_8d_db = pd.read_excel(
            EXCEL_MASTER_PATH,
            sheet_name=SHEET_8D_DB,
            header=header_row
        )

except Exception as e:
    st.sidebar.warning(
        "No se pudo leer la hoja de base de datos de informes 8D.\n"
        f"Detalle técnico: {e}"
    )
    df_8d_db = pd.DataFrame()


# ─────────────────────────────────────────────
# Selección y generación del informe
# ─────────────────────────────────────────────
if not df_8d_db.empty:
    col_pedido_8d = _find_col_any(
        df_8d_db,
        ["Pedido", "pedido", "nº pedido", "no pedido", "numero pedido", "n pedido"],
    ) or "Pedido"

    if col_pedido_8d not in df_8d_db.columns:
        st.sidebar.warning(
            "La base de datos 8D no contiene una columna de 'Pedido'. "
            "No se puede seleccionar un informe por número de Pedido."
        )
    else:
        opciones_pedido = (
            df_8d_db[col_pedido_8d]
            .dropna()
            .astype(str)
            .unique()
            .tolist()
        )
        opciones_pedido = sorted(opciones_pedido)

        if not opciones_pedido:
            st.sidebar.info(
                "La base BDCAPAINFORME8D no tiene registros con Nº de Pedido."
            )
        else:
            pedido_sel = st.sidebar.selectbox(
                "Selecciona Nº de Pedido",
                opciones_pedido,
                key="sel_pedido_8d",
            )

            if st.sidebar.button(
                "📄 Generar informe 8D (Word) para este Pedido",
                key="btn_gen_8d_word_pedido",
                help="Genera un informe 8D en Word basándose en el registro coincidente.",
            ):
                mask = df_8d_db[col_pedido_8d].astype(str) == str(pedido_sel)
                df_sel = df_8d_db[mask]

                if df_sel.empty:
                    st.sidebar.warning(
                        f"No se encontró ningún registro 8D con Pedido = {pedido_sel}."
                    )
                else:
                    reg = df_sel.iloc[-1]  # último informe registrado para ese pedido
                    buffer_docx = crear_doc_informe_8d(reg)

                    st.sidebar.download_button(
                        label="⬇️ Descargar informe 8D (Word)",
                        data=buffer_docx,
                        file_name=f"Informe8D_Pedido_{pedido_sel}.docx",
                        mime=(
                            "application/vnd.openxmlformats-"
                            "officedocument.wordprocessingml.document"
                        ),
                        key="dl_informe_8d_word_pedido",
                    )

else:
    st.sidebar.info(
        "Todavía no hay registros en BDCAPAINFORME8D para generar informes 8D."
    )

# ──────────────────────────────────────────────────────────────
# 6. ENCABEZADO PRINCIPAL (TARJETA DEGRADADA)
# ──────────────────────────────────────────────────────────────

st.markdown(
    """
<div style="
    background: linear-gradient(90deg, #0ea5e9 0%, #2563eb 45%, #1d4ed8 100%);
    border-radius: 18px;
    padding: 18px 22px 16px 22px;
    box-shadow: 0 20px 45px rgba(15, 23, 42, 0.7);
    color: #e5f2ff;
    border: 1px solid rgba(59, 130, 246, 0.7);
    margin-top: 14px;
">
  <div style="font-size:0.80rem;margin-bottom:4px;font-weight:600;opacity:0.9;">
    Sistema de Gestión Operativa de Calidad · ECOcero
  </div>
  <h1 style="margin:0; font-weight:800; font-size:1.6rem;">
    Cuadro de mando para control y seguimiento de incidencias
  </h1>
  <p style="margin:3px 0 0 0; color:#dbeafe; font-size:0.90rem;">
    Visor operativo para CAPA e informes 8D, conectado a la base de datos SQLite del sistema.
  </p>
</div>
""",
    unsafe_allow_html=True,
)

st.markdown("")

# ──────────────────────────────────────────────────────────────
# 7. KPIs GLOBALES INTELIGENTES (CAPA)
# ──────────────────────────────────────────────────────────────

total_capa = int(len(df_view_search))
total_8d_vista = int(len(df_8d_view))

# Abiertas / cerradas
abiertas = cerradas = 0
if col_estatus and not df_view_search.empty:
    est_vals = df_view_search[col_estatus].astype(str).str.lower()
    cerradas = int(est_vals.str.contains("cerr", na=False).sum())
    abiertas = total_capa - cerradas

# Costes aproximados
def _to_num(s) -> float:
    try:
        return float(str(s).replace(".", "").replace(",", "."))
    except Exception:
        return 0.0

cost_total = 0.0
if not df_view_search.empty and (col_cost_rep or col_cost_dev):
    cost_rep = df_view_search[col_cost_rep].map(_to_num) if col_cost_rep else 0.0
    cost_dev = df_view_search[col_cost_dev].map(_to_num) if col_cost_dev else 0.0
    cost_total = float((cost_rep + cost_dev).sum())

# SLA calculado sobre la vista actual
sla_ok, sla_ko, sla_series = compute_sla(df_view_search, sla_mode, sla_target)
sla_total_eval = sla_ok + sla_ko
sla_cumpl = _pct(sla_ok, sla_total_eval) if sla_total_eval > 0 else 0.0

# Lead time medio / mediano
lt_mean, lt_med = compute_closure_time(df_view_search)

# KPIs sobre acciones (contención / correctiva)
if not df_view_search.empty and col_acc_cont:
    acciones_cont = int(
        df_view_search[col_acc_cont]
        .astype(str)
        .str.strip()
        .replace({"": None, "nan": None})
        .notna()
        .sum()
    )
else:
    acciones_cont = 0

if not df_view_search.empty and col_acc_corr:
    acciones_corr = int(
        df_view_search[col_acc_corr]
        .astype(str)
        .str.strip()
        .replace({"": None, "nan": None})
        .notna()
        .sum()
    )
else:
    acciones_corr = 0

acciones_totales = acciones_cont + acciones_corr

# Reincidencias globales
r30_global, r60_global, df_re_global = compute_reincidences(
    df_view_search, re_key_a, re_key_b
)

# Indicador global de riesgo (0-100) + nivel tipo semáforo
if total_capa > 0:
    backlog_pct = _pct(abiertas, total_capa)   # más abiertas ⇒ más riesgo
else:
    backlog_pct = 0.0

sla_penalty = 100.0 - sla_cumpl               # menor SLA ⇒ más riesgo
rec_penalty = min(r60_global * 12.0, 100.0)   # cada reincidencia 60d suma riesgo

risk_global_score = round(
    0.40 * backlog_pct +    # peso fuerte al % de CAPA abiertas
    0.40 * sla_penalty +    # peso fuerte al incumplimiento de SLA
    0.20 * rec_penalty,     # peso moderado a reincidencias
    1,
)

if risk_global_score >= 75:
    risk_global_level = "ALTO (Rojo)"
elif risk_global_score >= 50:
    risk_global_level = "MEDIO (Ámbar)"
else:
    risk_global_level = "BAJO (Verde)"

# ── Textos formateados para la tarjeta ────────────────────────
pct_abiertas_txt = f"{_pct(abiertas, total_capa)}% abiertas" if total_capa > 0 else "N/D"
pct_cont_txt = f"{_pct(acciones_cont, total_capa)}% de CAPA" if total_capa > 0 else "N/D"
pct_corr_txt = f"{_pct(acciones_corr, total_capa)}% de CAPA" if total_capa > 0 else "N/D"
sla_cumpl_txt = f"{sla_cumpl:.1f}%"
lt_mean_txt = f"{lt_mean:.1f} d (media)" if lt_mean is not None else "N/D"
lt_med_txt = f"{lt_med:.1f} d" if lt_med is not None else "N/D"
cost_txt = (
    f"{cost_total:,.2f} €"
    .replace(",", "X")
    .replace(".", ",")
    .replace("X", ".")
)

# ── Tarjeta visual única para el resumen de KPIs ──────────────
st.markdown(
    f"""
<div class="eco-kpi-card">
  <div class="eco-kpi-grid">
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">CAPA totales (vista filtrada + búsqueda)</div>
      <div class="eco-kpi-value">{total_capa}</div>
    </div>
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">CAPA abiertas</div>
      <div class="eco-kpi-value">{abiertas}</div>
      <div class="eco-kpi-chip">{pct_abiertas_txt}</div>
    </div>
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">Cumplimiento SLA</div>
      <div class="eco-kpi-value">{sla_cumpl_txt}</div>
      <div class="eco-kpi-chip">Objetivo: {sla_target:.0f} {sla_mode}</div>
    </div>
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">Reincidencias (30 / 60 días)</div>
      <div class="eco-kpi-value">{r30_global} / {r60_global}</div>
      <div class="eco-kpi-chip">RC30 / RC60 (todas las claves)</div>
    </div>
  </div>

  <div class="eco-kpi-grid" style="margin-top:0.8rem;">
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">Acciones de contención (n≠0)</div>
      <div class="eco-kpi-value">{acciones_cont}</div>
      <div class="eco-kpi-chip">{pct_cont_txt}</div>
    </div>
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">Acciones correctivas (n≠0)</div>
      <div class="eco-kpi-value">{acciones_corr}</div>
      <div class="eco-kpi-chip">{pct_corr_txt}</div>
    </div>
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">Cierre del plazo de entrega CAPA</div>
      <div class="eco-kpi-value">{lt_mean_txt}</div>
      <div class="eco-kpi-chip">Mediana: {lt_med_txt}</div>
    </div>
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">Coste total estimado</div>
      <div class="eco-kpi-value">{cost_txt}</div>
    </div>
  </div>

  <div class="eco-kpi-grid" style="margin-top:0.9rem;">
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">Indicador global de riesgo</div>
      <div class="eco-kpi-value">{risk_global_score:.1f} / 100</div>
      <div class="eco-kpi-chip eco-kpi-chip--green">
        {risk_global_level}
      </div>
    </div>
  </div>
</div>
""",
    unsafe_allow_html=True,
)

st.markdown("---")

# ──────────────────────────────────────────────────────────────
# 8. LECTURA AUTOMÁTICA GLOBAL (ESTILO AUDITORÍA)
# ──────────────────────────────────────────────────────────────

def lectura_auditoria_global() -> str:
    """Genera una lectura en lenguaje natural para el resumen global de KPIs."""
    if total_capa == 0:
        return (
            "No se encuentran CAPA visibles con los filtros y el buscador actual. "
            "Antes de interpretar tendencias o riesgos, es necesario cargar datos "
            "o ampliar los filtros."
        )

    texto = []

    # Apertura vs cierre
    texto.append(
        f"Se observan **{total_capa} CAPA** en la vista actual, de las cuales "
        f"**{abiertas} ({_pct(abiertas, total_capa)}%)** permanecen abiertas y "
        f"**{cerradas} ({_pct(cerradas, total_capa)}%)** han sido cerradas."
    )

    if abiertas > cerradas:
        texto.append(
            "El volumen de CAPA abiertas es superior al de cerradas, lo cual indica "
            "una **cola de trabajo pendiente** que conviene priorizar."
        )
    elif cerradas > abiertas:
        texto.append(
            "El número de CAPA cerradas supera al de abiertas, lo que sugiere una "
            "**buena capacidad de cierre** respecto a las aperturas registradas."
        )
    else:
        texto.append(
            "El número de CAPA abiertas y cerradas es similar, reflejando un "
            "**equilibrio** entre la generación y el cierre de incidencias."
        )

    # SLA
    if sla_total_eval > 0:
        if sla_cumpl >= 80:
            texto.append(
                f"El **cumplimiento de SLA** se sitúa en **{sla_cumpl:.1f}%**, "
                "un valor **sólido**, alineado con buenas prácticas (≥80%)."
            )
        elif sla_cumpl >= 60:
            texto.append(
                f"El cumplimiento de SLA alcanza **{sla_cumpl:.1f}%**. Es un "
                "resultado **aceptable**, pero con margen para mejorar tiempos "
                "de respuesta y cierre."
            )
        else:
            texto.append(
                f"El SLA se cumple sólo en **{sla_cumpl:.1f}%** de los casos, "
                "lo que representa un **riesgo operativo**. Se recomienda revisar "
                "cuellos de botella, recursos y prioridades."
            )
    else:
        texto.append(
            "No se ha podido calcular el indicador de SLA (falta de columnas de fecha o datos suficientes)."
        )

    # Lead time
    if lt_mean is not None:
        texto.append(
            f"El **tiempo medio de cierre** de una CAPA es de aproximadamente "
            f"**{lt_mean:.1f} días**, con una mediana de **{(lt_med or lt_mean):.1f} días**. "
            "Valores muy elevados pueden indicar retrasos en la implantación y verificación de acciones."
        )

    # Acciones
    if acciones_totales > 0:
        texto.append(
            f"Existen **{acciones_cont} CAPA con acción de contención** y "
            f"**{acciones_corr} con acción correctiva**. El porcentaje de CAPA "
            f"con acciones registradas respecto al total es de "
            f"**{_pct(acciones_totales, total_capa)}%**."
        )
    else:
        texto.append(
            "No se observan acciones de contención o correctivas registradas en la vista actual, "
            "lo cual limita la trazabilidad de la respuesta frente a las incidencias."
        )

    # Reincidencias
    if r30_global == 0 and r60_global == 0:
        texto.append(
            "No se detectan **reincidencias significativas** en ventanas de 30 y 60 días, "
            "lo cual sugiere una buena contención de los problemas detectados."
        )
    else:
        texto.append(
            f"Se han identificado **{r30_global} reincidencias a 30 días** y "
            f"**{r60_global} a 60 días**. Es recomendable revisar estas combinaciones "
            "de cliente/tipo/área para evaluar la **eficacia real** de las acciones correctivas."
        )

    if cost_total > 0:
        texto.append(
            f"El **coste total estimado** asociado a las CAPA visibles se sitúa en "
            f"aproximadamente **{cost_total:,.2f} €** (suma de reposiciones y devoluciones)."
        )
# Indicador global de riesgo (semáforo)
    texto.append(
        f"Combinando backlog de CAPA abiertas, cumplimiento de SLA y reincidencias a 60 días, "
        f"el sistema estima un **indicador global de riesgo** de **{risk_global_score:.1f}/100**, "
        f"clasificado como **{risk_global_level}** en modo semáforo "
        "(verde = riesgo bajo, ámbar = riesgo medio, rojo = riesgo alto)."
    )
    return " ".join(texto)


with st.expander("🧾 Lectura automática global (estilo auditoría)", expanded=False):
    st.markdown(f"<div class='eco-audit'>{lectura_auditoria_global()}</div>", unsafe_allow_html=True)

st.markdown("---")

# ──────────────────────────────────────────────────────────────
# 9. PESTAÑAS PRINCIPALES DEL DASHBOARD (ESQUELETO ORGANIZADO)
# ──────────────────────────────────────────────────────────────

tab_resumen, tab_tablas, tab_incidencias, tab_sla_reincidencias, tab_acciones_eficacia, tab_indicadores, tab_costos, tab_tendencias, tab_plan_ia = st.tabs(
    [
        "📌 Resumen general",
        "📋 Tablas",
        "📈 Incidencias",
        "⏱️ SLA & 🔁 Reincidencias",
        "🧩 Acciones & ✅ Eficacia",
        "📊 Indicadores",
        "💸 Costos",
        "📈 Tendencias & pronóstico",
        "🌐 plan de acciones (IA)",
    ]
)

# ──────────────────────────────────────────────────────────────
# TAB 1 – Resumen general · Informe ejecutivo + técnico + sugerencias
# ──────────────────────────────────────────────────────────────
with tab_resumen:
    st.subheader("📌 Resumen general")

    # Estado para guardar último informe generado y gráficos asociados
    if "informe_resumen_texto" not in st.session_state:
        st.session_state["informe_resumen_texto"] = ""
    if "informe_resumen_grafs" not in st.session_state:
        st.session_state["informe_resumen_grafs"] = []

    if total_capa == 0:
        st.info(
            "En este momento el sistema no tiene registros CAPA cargados en la vista actual. "
            "Una vez se incorporen incidencias, aquí podrás generar un informe ejecutivo y técnico "
            "del desempeño global del sistema."
        )


    else:
        # ------------------------------------------------------
        # 1) Selección de periodo + configuración de informe
        # ------------------------------------------------------
        col_f1, col_f2 = st.columns(2)
        with col_f1:
            fecha_ini_inf = st.date_input(
                "📅 Fecha inicio del periodo del informe",
                value=None,
                format="DD/MM/YYYY",
            )
        with col_f2:
            fecha_fin_inf = st.date_input(
                "📅 Fecha fin del periodo del informe",
                value=None,
                format="DD/MM/YYYY",
            )

        # Texto descriptivo del periodo
        if fecha_ini_inf and fecha_fin_inf:
            periodo_txt = (
                f"entre el {fecha_ini_inf.strftime('%d/%m/%Y')} "
                f"y el {fecha_fin_inf.strftime('%d/%m/%Y')}"
            )
        else:
            periodo_txt = "según los filtros actuales del dashboard"

        st.markdown("---")

        st.markdown("##### 🎯 Configuración del informe")

        tipo_informe = st.radio(
            "¿Qué nivel de detalle quieres visualizar en pantalla?",
            options=["Ejecutivo (resumen)", "Ejecutivo + Técnico completo"],
            index=1,
            horizontal=True,
        )

        graficos_seleccionados = st.multiselect(
            "Si luego vas a anexar gráficos al informe, marca aquí cuáles te interesa incluir:",
            options=[
                "Pareto de defectos (Incidencias)",
                "Tendencia mensual de incidencias",
                "Cumplimiento global de SLA",
                "Distribución de reincidencias 30 / 60 días",
                "Costes totales por tipo (reposiciones vs devoluciones)",
                "Costes por tipo de incidencia / defecto",
                "Evolución de costes en el tiempo",
            ],
            default=[
                "Pareto de defectos (Incidencias)",
                "Cumplimiento global de SLA",
                "Distribución de reincidencias 30 / 60 días",
            ],
        )

        st.markdown("---")

        # ------------------------------------------------------
        # 2) Funciones de generación de texto
        # ------------------------------------------------------
        def lectura_ejecutiva_global(periodo: str) -> str:
            textos = []
            textos.append(
                f"Este informe resume el estado del sistema de gestión de incidencias CAPA "
                f"{periodo}, utilizando la vista y filtros activos en el dashboard."
            )

            backlog_pct = _pct(abiertas, total_capa) if total_capa > 0 else 0.0
            textos.append(
                f"Actualmente se gestionan **{total_capa} CAPA**, de las cuales "
                f"**{abiertas}** permanecen **abiertas** "
                f"({backlog_pct}% del total) y **{cerradas}** figuran como **cerradas**."
            )

            # SLA
            sla_total_eval = sla_ok + sla_ko
            if sla_total_eval > 0:
                textos.append(
                    f"El **cumplimiento global de SLA** se sitúa en **{sla_cumpl:.1f}%**, "
                    f"tomando como referencia un objetivo de **{sla_target:.0f} {sla_mode}** "
                    "para la resolución de las incidencias."
                )
            else:
                textos.append(
                    "Por el momento no hay suficientes registros con fechas de apertura y cierre "
                    "para calcular un indicador de SLA representativo."
                )

            # Reincidencias
            textos.append(
                f"En términos de recurrencia, se observan **{r30_global} reincidencias a 30 días** "
                f"y **{r60_global} a 60 días**, lo que permite evaluar la estabilidad de las "
                "soluciones implantadas."
            )

            # Acciones
            textos.append(
                f"Desde la perspectiva de las acciones, **{acciones_cont} CAPA** cuentan con "
                "alguna **acción de contención** y "
                f"**{acciones_corr} CAPA** disponen de al menos una **acción correctiva**."
            )

            # Lead time
            if lt_mean is not None and lt_med is not None:
                textos.append(
                    f"El **tiempo promedio de cierre** es de **{lt_mean:.1f} días**, "
                    f"con una mediana de **{lt_med:.1f} días**, lo que ofrece una referencia "
                    "clara sobre la velocidad real de resolución."
                )

            # Costes
            if cost_total > 0:
                textos.append(
                    f"El **coste total estimado** asociado a las incidencias asciende a "
                    f"aproximadamente **{cost_total:,.2f} €**."
                )

            # Riesgo global
            textos.append(
                f"Integrando backlog, SLA y reincidencias, el **indicador global de riesgo** "
                f"se sitúa en **{risk_global_score:.1f} / 100**, clasificado como "
                f"**{risk_global_level}**."
            )

            return " ".join(textos)

        def lectura_tecnica_detallada(periodo: str) -> str:
            textos = []
            textos.append(
                f"Desde un punto de vista técnico, el presente informe se basa en los registros "
                f"CAPA visibles {periodo}, tras aplicar los filtros seleccionados en el dashboard."
            )

            # Detalle SLA
            sla_total_eval = sla_ok + sla_ko
            if sla_total_eval > 0:
                textos.append(
                    f"El indicador de SLA se calcula como el cociente entre las CAPA cerradas "
                    "dentro del tiempo objetivo y el total de CAPA con fechas válidas de "
                    "apertura y cierre. En este caso, se evaluaron "
                    f"**{sla_total_eval} incidencias**, con un cumplimiento de "
                    f"**{sla_cumpl:.1f}%**."
                )

            # Reincidencias
            textos.append(
                f"Las métricas de reincidencia **RC30** y **RC60** agrupan las incidencias por "
                "su clave de reincidencia y contabilizan cuántas reaparecen dentro de ventanas "
                "temporales de 30 y 60 días, respectivamente. Actualmente se observan "
                f"**{r30_global} casos RC30** y **{r60_global} casos RC60**."
            )

            # Acciones
            textos.append(
                f"Los indicadores de acciones se construyen revisando si los campos de acciones "
                "de contención y correctivas contienen información no vacía. De las "
                f"**{total_capa} CAPA** analizadas, **{acciones_cont}** presentan "
                "acciones de contención y "
                f"**{acciones_corr}** tienen acciones correctivas definidas."
            )

            # Costes
            if cost_total > 0:
                textos.append(
                    "Los indicadores económicos utilizan los campos de costes de reposición y "
                    "devolución, transformados a valores numéricos. La suma de estos importes "
                    f"da lugar a un **coste total estimado de {cost_total:,.2f} €** para el "
                    "conjunto de la vista."
                )

            # Riesgo global
            textos.append(
                "El indicador global de riesgo se construye combinando tres componentes: "
                "**porcentaje de CAPA abiertas (backlog)**, **incumplimiento de SLA** y "
                "**reincidencias a 60 días**. Cada componente se normaliza y pondera para "
                "obtener un índice entre 0 y 100, donde valores más altos representan "
                "mayor tensión en el sistema de calidad."
            )

            return " ".join(textos)

        def generar_sugerencias_mejora() -> str:
            textos = []
            textos.append(
                "A partir de los indicadores analizados, se proponen las siguientes "
                "**líneas de mejora** y acciones recomendadas:"
            )

            # Según riesgo
            if "ALTO" in risk_global_level:
                textos.append(
                    "- **Priorizar la reducción del backlog** de CAPA abiertas, definiendo "
                    "compromisos claros de cierre por responsable y plazo.\n"
                )
            elif "MEDIO" in risk_global_level:
                textos.append(
                    "- **Consolidar el cierre oportuno** de las CAPA abiertas y reforzar la "
                    "revisión de las acciones correctivas con mayor impacto.\n"
                )
            else:
                textos.append(
                    "- **Mantener la disciplina actual** en el registro y seguimiento de CAPA, "
                    "reforzando las buenas prácticas ya implantadas.\n"
                )

            # SLA
            if sla_ok + sla_ko > 0 and sla_cumpl < 90:
                textos.append(
                    "- Revisar los **cuellos de botella en el flujo de resolución** para "
                    "aumentar el cumplimiento de SLA, especialmente en las áreas o clientes "
                    "con mayor volumen de incidencias fuera de plazo.\n"
                )

            # Reincidencias
            if r60_global > 0:
                textos.append(
                    "- Analizar en detalle las **incidencias reincidentes (RC30 / RC60)**, "
                    "verificando la calidad de la identificación de causa raíz y la "
                    "implantación/verificación de las acciones correctivas.\n"
                )

            # Acciones
            if acciones_corr == 0 or acciones_cont == 0:
                textos.append(
                    "- Incrementar la **formalización de acciones de contención y correctivas** "
                    "para cada CAPA relevante, evitando cierres sin plan de acción documentado.\n"
                )

            # Costes
            if cost_total > 0:
                textos.append(
                    "- Utilizar los **costes por tipo de incidencia y por cliente/proceso** "
                    "para priorizar proyectos de mejora en aquellos problemas que generan "
                    "mayor impacto económico.\n"
                )

            textos.append(
                "Estas recomendaciones deben revisarse en comité de calidad, asignando "
                "responsables, plazos e indicadores de seguimiento para cada acción."
            )

            return " ".join(textos)

        def construir_seccion_graficos(lista_graficos) -> str:
            if not lista_graficos:
                return (
                    "En este informe no se ha seleccionado ningún gráfico específico. "
                    "Se recomienda, no obstante, apoyarse en las pestañas del dashboard "
                    "para visualizar la evolución de incidencias, SLA y costes."
                )

            textos = []
            textos.append(
                "Para complementar este informe textual, se recomienda anexar los "
                "**gráficos generados en el dashboard** correspondientes a:"
            )
            for g in lista_graficos:
                textos.append(f"- {g}")
            textos.append(
                "Estos gráficos deben extraerse desde las pestañas respectivas del dashboard "
                "y adjuntarse como soporte visual en presentaciones o reportes formales."
            )
            return "\n".join(textos)

        # ------------------------------------------------------
        # 3) Generación del informe (texto) y guardado en sesión
        # ------------------------------------------------------
        if st.button("📝 Generar informe ejecutivo + técnico"):
            texto_exec = lectura_ejecutiva_global(periodo_txt)
            texto_tec = lectura_tecnica_detallada(periodo_txt)
            texto_sug = generar_sugerencias_mejora()
            texto_grafs = construir_seccion_graficos(graficos_seleccionados)

            if tipo_informe.startswith("Ejecutivo (resumen)"):
                cuerpo_informe = "\n\n".join(
                    [
                        "RESUMEN EJECUTIVO",
                        texto_exec,
                        "\nSUGERENCIAS DE MEJORA",
                        texto_sug,
                        "\nGRÁFICOS RECOMENDADOS",
                        texto_grafs,
                    ]
                )
            else:
                cuerpo_informe = "\n\n".join(
                    [
                        "RESUMEN EJECUTIVO",
                        texto_exec,
                        "\nINFORME TÉCNICO DETALLADO",
                        texto_tec,
                        "\nSUGERENCIAS DE MEJORA",
                        texto_sug,
                        "\nGRÁFICOS RECOMENDADOS",
                        texto_grafs,
                    ]
                )

            # Guardamos en sesión para previsualizar y descargar como DOCX
            st.session_state["informe_resumen_texto"] = cuerpo_informe
            st.session_state["informe_resumen_grafs"] = graficos_seleccionados
            st.success("Informe generado. Revisa la previsualización y descarga el Word más abajo.")

        # ------------------------------------------------------
        # 4) Previsualización + descarga en Word (DOCX)
        # ------------------------------------------------------
        cuerpo_informe = st.session_state.get("informe_resumen_texto", "")
        grafs_para_docx = st.session_state.get("informe_resumen_grafs", [])

        if cuerpo_informe:
            st.markdown("##### 👁️ Previsualización del informe")
            st.markdown(
                f"<div class='eco-audit'>{cuerpo_informe.replace(chr(10), '<br>')}</div>",
                unsafe_allow_html=True,
            )

            # Construir DOCX en memoria con texto + gráficos seleccionados
            docx_buffer = construir_docx_informe(cuerpo_informe, grafs_para_docx)

            st.download_button(
                "⬇️ Descargar informe en Word (DOCX)",
                data=docx_buffer,
                file_name="informe_resumen_capa.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        else:
            st.info("Genera primero el informe para poder previsualizarlo y descargarlo en Word.")

# ──────────────────────────────────────────────────────────────
# PESTAÑA: 📋 TABLAS (solo datos, sin KPIs ni gráficos)
# ──────────────────────────────────────────────────────────────
with tab_tablas:
    st.subheader("📋 Tablas de datos (vista filtrada)")

    # Aseguramos que existan los dataframes de vista
    df_capa_tab = df_view_search if "df_view_search" in globals() else df_view
    df_8d_tab = df_8d_view if "df_8d_view" in globals() else None

    if (df_capa_tab is None or df_capa_tab.empty) and (df_8d_tab is None or df_8d_tab.empty):
        st.info(
            "Por ahora no hay registros para mostrar en las tablas. "
            "Carga o filtra datos en la pestaña de incidencias para ver información aquí."
        )
    else:
        # Subpestañas internas solo para organizar tablas
        t_capa, t_8d = st.tabs(
            [
                "📄 CAPA (BDCAPA)",
                "📝 Informes 8D (BDCAPAINFORME8D)",
            ]
        )

        # ------------------------------------------------------
        # TABLA CAPA
        # ------------------------------------------------------
        with t_capa:
            st.markdown("#### 📄 Tabla CAPA (vista actual)")

            if df_capa_tab is None or df_capa_tab.empty:
                st.info("No hay registros CAPA en la vista actual.")
            else:
                st.caption(
                    "Cada fila representa un registro CAPA según los filtros seleccionados "
                    "en la barra lateral."
                )
                st.dataframe(df_capa_tab, width='stretch')

                # Descarga de la tabla CAPA
                csv_capa = df_capa_tab.to_csv(index=False).encode("utf-8-sig")
                st.download_button(
                    label="⬇️ Descargar tabla CAPA (CSV)",
                    data=csv_capa,
                    file_name="CAPA_vista_filtrada.csv",
                    mime="text/csv",
                    key="dl_capa_tablas",
                )

        # ------------------------------------------------------
        # TABLA 8D
        # ------------------------------------------------------
        with t_8d:
            st.markdown("#### 📝 Tabla Informes 8D (vista actual)")

            if df_8d_tab is None or df_8d_tab.empty:
                st.info(
                    "No hay registros 8D en la vista actual o aún no se ha cargado la hoja "
                    "'BDCAPAINFORME8D' del Excel."
                )
            else:
                st.caption(
                    "Cada fila representa un informe 8D asociado a las incidencias, según la "
                    "vista filtrada actual."
                )
                st.dataframe(df_8d_tab, width='stretch')

                # Descarga de la tabla 8D
                csv_8d = df_8d_tab.to_csv(index=False).encode("utf-8-sig")
                st.download_button(
                    label="⬇️ Descargar tabla 8D (CSV)",
                    data=csv_8d,
                    file_name="Informes8D_vista_filtrada.csv",
                    mime="text/csv",
                    key="dl_8d_tablas",
                )
# ──────────────────────────────────────────────────────────────
# TAB 2 – Incidencias · Catálogo de defectos & Pareto
# ──────────────────────────────────────────────────────────────
with tab_incidencias:
    st.markdown("#### 📈 Incidencias · Catálogo de defectos & Pareto")

    # Usamos la vista ya filtrada + buscador global si existe
    df_def = df_view_search if "df_view_search" in globals() else df_view
    if df_def is None:
        df_def = pd.DataFrame()

    # Si no hay datos, avisamos y salimos
    if df_def.empty:
        st.info("No hay registros para analizar con los filtros actuales.")
    else:
        # Intentamos detectar la columna de defecto / tipo incidencia
        col_defecto = _find_col_any(
            df_def,
            [
                "defecto",
                "defectos",
                "tipo incidencia",
                "incidencia",
                "motivo",
                "descripción detallada de la incidencia",
                "descripcion detallada de la incidencia",
            ],
        )

        # Si no detectamos columna de defecto, informamos
        if not col_defecto:
            st.info(
                "No se encontró una columna clara de defectos/tipo de incidencia. "
                "Revisa el diccionario de datos para confirmar los nombres."
            )
        else:
            st.markdown("### 📉 Pareto de defectos (vista filtrada)")

            # Normalizamos la columna de defecto
            serie_def = (
                df_def[col_defecto]
                .astype(str)
                .str.strip()
                .replace({"nan": "", "None": ""})
            )

            # Quitamos vacíos
            serie_def = serie_def[serie_def != ""]

            if serie_def.empty:
                st.info(
                    f"La columna detectada como defecto (**{col_defecto}**) "
                    "no contiene valores significativos en la vista actual."
                )
            else:
                # Conteo por tipo de defecto
                vc_def = serie_def.value_counts()
                total_def = int(vc_def.sum())

                df_pareto = vc_def.reset_index()
                df_pareto.columns = ["Defecto", "Recuento"]

                # % sobre el total (evitando división por cero)
                if total_def > 0:
                    df_pareto["% sobre total"] = df_pareto["Recuento"].apply(
                        lambda x: _pct(x, total_def)
                    )
                else:
                    df_pareto["% sobre total"] = 0.0

                # % acumulado típico de Pareto
                df_pareto = df_pareto.sort_values("Recuento", ascending=False)
                df_pareto["% acumulado"] = df_pareto["% sobre total"].cumsum()

                # Tabla Pareto
                st.dataframe(df_pareto, width='stretch')

                # Gráfico de barras del Pareto
                try:
                    fig = px.bar(
                        df_pareto,
                        x="Defecto",
                        y="Recuento",
                        title="Pareto de defectos",
                        text="Recuento",
                    )
                    fig.update_traces(textposition="outside", cliponaxis=False)
                    fig.update_layout(
                        margin=dict(l=10, r=10, t=40, b=140),
                        height=420,
                        xaxis_title=None,
                        yaxis_title="Nº incidencias",
                    )
                    fig.update_xaxes(tickangle=-30, automargin=True)

                    # 🔹 Registramos este gráfico para el informe en Word
                    caption_pareto = (
                        "Figura 1. Pareto de defectos/tipos de incidencia en la vista filtrada. "
                        "Cada barra representa el número de incidencias por defecto; el análisis "
                        "permite identificar los tipos de fallo más frecuentes y priorizar acciones."
                    )
                    if "FIGS_FOR_REPORT" in globals():
                        FIGS_FOR_REPORT["pareto_defectos"] = {
                            "fig": fig,
                            "caption": caption_pareto,
                            "section": "Incidencias – Pareto de defectos",
                            "created_at": pd.Timestamp.now(),
                        }

                    st.plotly_chart(fig, width='stretch')

                except Exception as e:
                    st.warning(f"Error al renderizar gráfico de Pareto: {e}")

                # ---------------- Lectura automática del Pareto ----------------
                def lectura_auditoria_defectos() -> str:
                    textos = []

                    textos.append(
                        f"En la vista actual se analizan **{total_def} incidencias** "
                        f"con un catálogo de defectos/tipos que incluye al menos "
                        f"**{len(vc_def)} categorías distintas**."
                    )

                    # Principal defecto
                    if not vc_def.empty:
                        d_principal = vc_def.index[0]
                        n_principal = int(vc_def.iloc[0])
                        p_principal = _pct(n_principal, total_def)
                        textos.append(
                            f"El defecto o tipo de incidencia más frecuente es "
                            f"**{d_principal}**, con **{n_principal} casos "
                            f"({p_principal}% del total)**."
                        )

                    # Defectos que acumulan ~80% del total (regla de Pareto)
                    top_80 = df_pareto[df_pareto["% acumulado"] <= 80.0]
                    if not top_80.empty:
                        categorias_80 = top_80["Defecto"].tolist()
                        textos.append(
                            "Siguiendo el principio de Pareto, aproximadamente el **80% de las "
                            "incidencias** se concentra en los siguientes defectos clave: "
                            + ", ".join(f"**{c}**" for c in categorias_80)
                            + "."
                        )
                    else:
                        textos.append(
                            "No se identifica un conjunto reducido de defectos que concentre el 80% "
                            "de las incidencias; esto sugiere una distribución más dispersa de los fallos."
                        )

                    textos.append(
                        "Se recomienda utilizar este Pareto como base para priorizar las acciones de "
                        "mejora, enfocando primero en los defectos con mayor peso relativo y "
                        "revisando la eficacia de las acciones implantadas sobre ellos."
                    )

                    return " ".join(textos)

                with st.expander(
                    "🧾 Lectura automática de defectos (estilo auditoría)", expanded=False
                ):
                    st.markdown(
                        f"<div class='eco-audit'>{lectura_auditoria_defectos()}</div>",
                        unsafe_allow_html=True,
                    )

# ──────────────────────────────────────────────────────────────
# TAB 3 – SLA & Reincidencias
# ──────────────────────────────────────────────────────────────
with tab_sla_reincidencias:
    st.markdown("#### ⏱️ SLA & Reincidencias")

    # Si no hay CAPA en la vista actual, mensaje y salimos
    if total_capa == 0:
        st.info(
            "Por ahora no hay registros CAPA en la vista actual. "
            "Cuando se carguen incidencias, esta pestaña mostrará "
            "los análisis de SLA y de reincidencias."
        )
    else:
        # Usamos la vista filtrada + buscador global
        df_sla = df_view_search if "df_view_search" in globals() else df_view
        if df_sla is None:
            df_sla = pd.DataFrame()

        # ──────────────────────────────────────────────────────
        # BLOQUE A · SLA (cumplimiento de tiempos)
        # ──────────────────────────────────────────────────────
        st.markdown("### ⏱️ Análisis de SLA (tiempos de respuesta)")

        sla_total_eval = sla_ok + sla_ko
        sla_txt_total = (
            f"{sla_total_eval} CAPA evaluadas" if sla_total_eval > 0 else "Sin datos suficientes"
        )

        # Tarjeta resumen SLA
        st.markdown(
            f"""
<div class="eco-kpi-card">
  <div class="eco-kpi-grid">
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">Cumplimiento global de SLA</div>
      <div class="eco-kpi-value">{sla_cumpl:.1f}%</div>
      <div class="eco-kpi-chip">
        Objetivo: {sla_target:.0f} {sla_mode}
      </div>
    </div>
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">CAPA dentro del SLA</div>
      <div class="eco-kpi-value">{sla_ok}</div>
      <div class="eco-kpi-chip">
        En plazo
      </div>
    </div>
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">CAPA fuera del SLA</div>
      <div class="eco-kpi-value">{sla_ko}</div>
      <div class="eco-kpi-chip">
        Fuera de plazo
      </div>
    </div>
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">Base de cálculo</div>
      <div class="eco-kpi-value">{sla_txt_total}</div>
    </div>
  </div>
</div>
            """,
            unsafe_allow_html=True,
        )

        # Gráfico simple de barras SLA (dentro vs fuera)
        if sla_total_eval > 0:
            try:
                df_sla_plot = pd.DataFrame(
                    {
                        "Estado SLA": ["Dentro SLA", "Fuera SLA"],
                        "CAPA": [sla_ok, sla_ko],
                    }
                )
                fig_sla = px.bar(
                    df_sla_plot,
                    x="Estado SLA",
                    y="CAPA",
                    title="Distribución de CAPA dentro / fuera de SLA",
                    text="CAPA",
                )
                fig_sla.update_traces(textposition="outside", cliponaxis=False)
                fig_sla.update_layout(
                    margin=dict(l=10, r=10, t=40, b=60),
                    height=380,
                    xaxis_title=None,
                    yaxis_title="Nº CAPA",
                )

                # 🔹 Registrar este gráfico para el informe Word
                caption_sla = (
                    "Figura X. Distribución de CAPA dentro y fuera del SLA objetivo. "
                    "Permite visualizar el equilibrio entre incidencias resueltas en plazo "
                    "y aquellas que superan el tiempo comprometido."
                )
                if "FIGS_FOR_REPORT" in globals():
                    FIGS_FOR_REPORT["sla_dentro_fuera"] = {
                        "fig": fig_sla,
                        "caption": caption_sla,
                        "section": "SLA – Distribución dentro / fuera de plazo",
                        "created_at": pd.Timestamp.now(),
                    }

                st.plotly_chart(fig_sla, width='stretch')

            except Exception as e:
                st.warning(f"Error al renderizar gráfico de SLA: {e}")

        # Lectura automática SLA
        def lectura_sla_resumen() -> str:
            if sla_total_eval == 0:
                return (
                    "Todavía no hay suficientes registros con fechas de apertura y cierre "
                    "para calcular un SLA robusto. Es importante completar estos campos "
                    "para poder evaluar la capacidad de respuesta del sistema."
                )

            textos = []
            textos.append(
                f"El **cumplimiento global de SLA** alcanza **{sla_cumpl:.1f}%** sobre "
                f"un total de **{sla_total_eval} CAPA** evaluadas. "
                f"De ellas, **{sla_ok}** se cerraron dentro del objetivo de "
                f"**{sla_target:.0f} {sla_mode}**, mientras que **{sla_ko}** "
                "superaron dicho umbral."
            )

            if sla_cumpl >= 90:
                textos.append(
                    "Este nivel de cumplimiento refleja una **muy buena capacidad de respuesta**, "
                    "coherente con un sistema de gestión maduro."
                )
            elif sla_cumpl >= 70:
                textos.append(
                    "El cumplimiento es **aceptable**, pero muestra margen de mejora, "
                    "especialmente en las CAPA que exceden el plazo objetivo."
                )
            else:
                textos.append(
                    "El nivel de cumplimiento es **bajo**, lo que indica que una parte "
                    "importante de las incidencias no se resuelven en el tiempo objetivo. "
                    "Es un foco prioritario para revisar recursos, flujos y cuellos de botella."
                )

            textos.append(
                "Técnicamente, este indicador se calcula como el porcentaje de CAPA cerradas "
                "dentro del tiempo objetivo sobre el total de CAPA con fechas válidas de "
                "apertura y cierre."
            )
            textos.append(
                "En las reuniones de seguimiento, este bloque de SLA debe utilizarse como "
                "referencia directa de la capacidad de respuesta del sistema y de la "
                "priorización de recursos."
            )
            return " ".join(textos)

        with st.expander("🧾 Lectura automática de SLA (estilo auditoría)", expanded=False):
            st.markdown(
                f"<div class='eco-audit'>{lectura_sla_resumen()}</div>",
                unsafe_allow_html=True,
            )

        st.markdown("---")

        # ──────────────────────────────────────────────────────
        # BLOQUE B · REINCIDENCIAS (RC30 / RC60)
        # ──────────────────────────────────────────────────────
        st.markdown("### 🔁 Análisis de reincidencias (RC30 / RC60)")

        # Tarjeta resumen reincidencias
        st.markdown(
            f"""
<div class="eco-kpi-card">
  <div class="eco-kpi-grid">
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">Reincidencias 30 días (RC30)</div>
      <div class="eco-kpi-value">{r30_global}</div>
    </div>
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">Reincidencias 60 días (RC60)</div>
      <div class="eco-kpi-value">{r60_global}</div>
    </div>
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">CAPA con acciones correctivas (base análisis)</div>
      <div class="eco-kpi-value">{acciones_corr}</div>
    </div>
  </div>
</div>
            """,
            unsafe_allow_html=True,
        )

        # Gráfico de barras RC30 / RC60
        try:
            df_reinc_plot = pd.DataFrame(
                {
                    "Ventana": ["RC30", "RC60"],
                    "Reincidencias": [r30_global, r60_global],
                }
            )
            fig_reinc = px.bar(
                df_reinc_plot,
                x="Ventana",
                y="Reincidencias",
                title="Reincidencias por ventana temporal (30 / 60 días)",
                text="Reincidencias",
            )
            fig_reinc.update_traces(textposition="outside", cliponaxis=False)
            fig_reinc.update_layout(
                margin=dict(l=10, r=10, t=40, b=60),
                height=380,
                xaxis_title=None,
                yaxis_title="Nº reincidencias",
            )

            # 🔹 Registrar este gráfico para el informe Word
            caption_reinc = (
                "Figura Y. Reincidencias a 30 y 60 días (RC30 / RC60). "
                "Permite comparar la recurrencia de los problemas en dos ventanas temporales "
                "y valorar la eficacia de las acciones correctivas."
            )
            if "FIGS_FOR_REPORT" in globals():
                FIGS_FOR_REPORT["reincidencias_rc30_rc60"] = {
                    "fig": fig_reinc,
                    "caption": caption_reinc,
                    "section": "Reincidencias – RC30 / RC60",
                    "created_at": pd.Timestamp.now(),
                }

            st.plotly_chart(fig_reinc, width='stretch')

        except Exception as e:
            st.warning(f"Error al renderizar gráfico de reincidencias: {e}")

        # Lectura automática reincidencias
        def lectura_reinc_resumen() -> str:
            textos = []
            textos.append(
                f"En la vista actual se identifican **{r30_global} reincidencias a 30 días** "
                f"y **{r60_global} reincidencias a 60 días**. Estas métricas permiten evaluar "
                "si los problemas tienden a reaparecer después de aplicar acciones correctivas."
            )

            if acciones_corr > 0:
                textos.append(
                    f"El número de reincidencias se interpreta siempre en relación con las "
                    f"**{acciones_corr} CAPA** que tienen acción correctiva registrada. "
                    "Un volumen alto de RC30/RC60 frente al número de acciones indica "
                    "posibles debilidades en la identificación de la causa raíz o en la "
                    "implantación/verificación de las acciones."
                )
            else:
                textos.append(
                    "Actualmente no hay acciones correctivas registradas; por ello, las "
                    "reincidencias deben interpretarse con cautela y se recomienda avanzar "
                    "en la formalización de acciones para cada CAPA relevante."
                )

            textos.append(
                "Técnicamente, RC30 y RC60 se calculan agrupando incidencias por su clave "
                "de reincidencia y midiendo cuántas vuelven a aparecer dentro de las "
                "ventanas temporales de 30 y 60 días."
            )
            textos.append(
                "En las reuniones de revisión, este bloque ayuda a distinguir entre problemas "
                "puntuales y problemas estructurales que reaparecen, priorizando la mejora en "
                "las áreas con mayor recurrencia."
            )
            return " ".join(textos)

        with st.expander(
            "🧾 Lectura automática de reincidencias (estilo auditoría)", expanded=False
        ):
            st.markdown(
                f"<div class='eco-audit'>{lectura_reinc_resumen()}</div>",
                unsafe_allow_html=True,
            )
# ──────────────────────────────────────────────────────────────
# TAB 4 – Acciones & Eficacia
# ──────────────────────────────────────────────────────────────
with tab_acciones_eficacia:
    st.markdown("#### 🧩 Acciones y eficacia de las CAPA")

    # Usamos la vista ya filtrada + buscador global si existe
    df_acc = df_view_search if "df_view_search" in globals() else df_view
    if df_acc is None:
        df_acc = pd.DataFrame()

    if df_acc.empty:
        st.info(
            "Por ahora no hay registros en la vista actual. "
            "Cuando existan incidencias, aquí verás un resumen de acciones de contención, "
            "correctivas y su lectura de eficacia."
        )
    else:
        # Volvemos a localizar columnas de acciones por seguridad
        col_acc_cont_loc = col_acc_cont
        col_acc_corr_loc = col_acc_corr

        if not col_acc_cont_loc and not col_acc_corr_loc:
            st.info(
                "No se han detectado columnas de acciones de contención/correctivas en los datos. "
                "Revisa el diccionario de datos para confirmar los nombres."
            )
        else:
            df_acc_local = df_acc.copy()

            # Normalizamos columnas (evitamos errores si no existen)
            if col_acc_cont_loc and col_acc_cont_loc in df_acc_local.columns:
                s_cont = (
                    df_acc_local[col_acc_cont_loc]
                    .astype(str)
                    .str.strip()
                    .replace({"": None, "nan": None, "None": None})
                )
            else:
                s_cont = pd.Series([None] * len(df_acc_local), index=df_acc_local.index)

            if col_acc_corr_loc and col_acc_corr_loc in df_acc_local.columns:
                s_corr = (
                    df_acc_local[col_acc_corr_loc]
                    .astype(str)
                    .str.strip()
                    .replace({"": None, "nan": None, "None": None})
                )
            else:
                s_corr = pd.Series([None] * len(df_acc_local), index=df_acc_local.index)

            # Flags de presencia de acciones
            mask_cont = s_cont.notna()
            mask_corr = s_corr.notna()
            mask_ambas = mask_cont & mask_corr
            mask_ninguna = ~mask_cont & ~mask_corr

            n_total = int(len(df_acc_local))
            n_cont = int(mask_cont.sum())
            n_corr = int(mask_corr.sum())
            n_ambas = int(mask_ambas.sum())
            n_ninguna = int(mask_ninguna.sum())

            # ─────────────────────────────────────────────
            # Tarjeta principal de KPIs de acciones
            # ─────────────────────────────────────────────
            st.markdown(
                f"""
<div class="eco-kpi-card">
<div class="eco-kpi-grid">
<div class="eco-kpi-item">
<div class="eco-kpi-label">CAPA totales (vista filtrada + búsqueda)</div>
<div class="eco-kpi-value">{n_total}</div>
</div>

<div class="eco-kpi-item">
<div class="eco-kpi-label">CAPA con acciones de contención</div>
<div class="eco-kpi-value">{n_cont}</div>
<div class="eco-kpi-chip">
{_pct(n_cont, n_total)}% del total
</div>
</div>

<div class="eco-kpi-item">
<div class="eco-kpi-label">CAPA con acciones correctivas</div>
<div class="eco-kpi-value">{n_corr}</div>
<div class="eco-kpi-chip">
{_pct(n_corr, n_total)}% del total
</div>
</div>

<div class="eco-kpi-item">
<div class="eco-kpi-label">CAPA con contención + correctiva</div>
<div class="eco-kpi-value">{n_ambas}</div>
<div class="eco-kpi-chip">
Cobertura completa
</div>
</div>

<div class="eco-kpi-item">
<div class="eco-kpi-label">CAPA sin acciones registradas</div>
<div class="eco-kpi-value">{n_ninguna}</div>
<div class="eco-kpi-chip">
{_pct(n_ninguna, n_total)}% del total
</div>
</div>
</div>
</div>
                """,
                unsafe_allow_html=True,
            )

            # ─────────────────────────────────────────────
            # Gráfico: distribución de tipos de cobertura
            # ─────────────────────────────────────────────
            tipos = [
                "Solo contención",
                "Solo correctiva",
                "Contención + correctiva",
                "Sin acciones",
            ]
            recuentos = [
                int((mask_cont & ~mask_corr).sum()),
                int((mask_corr & ~mask_cont).sum()),
                n_ambas,
                n_ninguna,
            ]

            df_cov = pd.DataFrame(
                {"Tipo de cobertura": tipos, "Recuento": recuentos}
            )

            st.markdown("##### 📊 Cobertura de acciones sobre las CAPA")
            try:
                fig_cov = px.bar(
                    df_cov,
                    x="Tipo de cobertura",
                    y="Recuento",
                    text="Recuento",
                    title="Distribución de CAPA según tipo de acciones registradas",
                )
                fig_cov.update_traces(textposition="outside", cliponaxis=False)
                fig_cov.update_layout(
                    margin=dict(l=10, r=10, t=60, b=80),
                    height=420,
                    xaxis_title=None,
                    yaxis_title="Nº de CAPA",
                )
                fig_cov.update_xaxes(tickangle=-20, automargin=True)

                # 🔹 Registrar este gráfico para el informe Word
                caption_cov = (
                    "Figura Z. Distribución de las CAPA según el tipo de acciones registradas "
                    "(solo contención, solo correctiva, ambas o sin acciones). "
                    "Permite evaluar el grado de formalización del tratamiento de incidencias."
                )
                if "FIGS_FOR_REPORT" in globals():
                    FIGS_FOR_REPORT["acciones_cobertura"] = {
                        "fig": fig_cov,
                        "caption": caption_cov,
                        "section": "Acciones y eficacia – Cobertura de acciones",
                        "created_at": pd.Timestamp.now(),
                    }

                st.plotly_chart(fig_cov, width='stretch')
            except Exception as e:
                st.warning(f"Error al renderizar el gráfico de cobertura de acciones: {e}")

            # ─────────────────────────────────────────────
            # Tabla de CAPA sin ninguna acción declarada
            # ─────────────────────────────────────────────
            if n_ninguna > 0:
                st.markdown("##### 🔍 CAPA sin acciones de contención ni correctivas")
                df_sin_acc = df_acc_local[mask_ninguna].copy()
                st.dataframe(df_sin_acc, width='stretch')

            # ─────────────────────────────────────────────
            # Lectura automática de acciones & eficacia
            # ─────────────────────────────────────────────
            def lectura_acciones_eficacia() -> str:
                textos = []

                textos.append(
                    f"En la vista actual se analizan **{n_total} CAPA**, de las cuales "
                    f"**{n_cont} ({_pct(n_cont, n_total)}%)** disponen de al menos una "
                    f"**acción de contención** y **{n_corr} ({_pct(n_corr, n_total)}%)** "
                    "incluyen alguna **acción correctiva**."
                )

                if n_ambas > 0:
                    textos.append(
                        f"Un subconjunto de **{n_ambas} CAPA** cuenta simultáneamente con "
                        "acciones de contención y correctivas, lo que indica un tratamiento "
                        "formal y completo del incidente."
                    )

                if n_ninguna > 0:
                    textos.append(
                        f"Se identifican **{n_ninguna} CAPA** sin acciones registradas. "
                        "Este grupo debería revisarse de forma prioritaria en comité de calidad, "
                        "ya que refleja incidencias sin un plan explícito de tratamiento."
                    )

                # Conexión con reincidencias y riesgo global (si existen)
                textos.append(
                    f"En paralelo, el sistema registra **{r30_global} reincidencias a 30 días** "
                    f"y **{r60_global} a 60 días**, lo que proporciona una medida indirecta "
                    "de la eficacia de las acciones implantadas. Valores elevados de RC30/RC60 "
                    "comparados con el volumen de acciones correctivas sugieren que es necesario "
                    "revisar la calidad de la identificación de la causa raíz y la verificación "
                    "de eficacia."
                )

                textos.append(
                    f"El **indicador global de riesgo** se sitúa en **{risk_global_score:.1f} / 100** "
                    f"({risk_global_level}), por lo que la combinación de backlog, SLA y "
                    "reincidencias debe utilizarse como contexto para priorizar qué CAPA y "
                    "qué acciones requieren mayor foco."
                )

                textos.append(
                    "Desde el punto de vista técnico, estos indicadores se obtienen revisando si "
                    "los campos de acciones contienen información no vacía para cada CAPA, y se "
                    "relacionan con los indicadores de reincidencia y SLA para evaluar su eficacia."
                )

                return " ".join(textos)

            with st.expander(
                "🧾 Lectura automática de acciones y eficacia (estilo auditoría)",
                expanded=False,
            ):
                st.markdown(
                    f"<div class='eco-audit'>{lectura_acciones_eficacia()}</div>",
                    unsafe_allow_html=True,
                )

# ──────────────────────────────────────────────────────────────
# PESTAÑA: 📊 INDICADORES (solo KPIs tipo tarjeta + lectura)
# ──────────────────────────────────────────────────────────────
with tab_indicadores:
    st.subheader("📊 Indicadores clave del sistema")

    if total_capa == 0:
        st.info(
            "Por ahora no hay registros CAPA en la vista actual. "
            "Cuando se carguen incidencias, esta pestaña mostrará los "
            "indicadores clave organizados por categoría, con su lectura automática."
        )
    else:
        # Sub-pestañas internas SOLO de indicadores (sin gráficos ni tablas)
        t_glob, t_sla, t_reinc, t_acc, t_risk, t_cost = st.tabs(
            [
                "🌐 Globales",
                "⏱️ SLA",
                "🔁 Reincidencias",
                "🧩 Acciones",
                "🚦 Riesgo & backlog",
                "💶 Costos",
            ]
        )

        # ------------------------------------------------------
        # 1) INDICADORES GLOBALES
        # ------------------------------------------------------
        with t_glob:
            st.markdown("### 🌐 Indicadores globales de CAPA")

            st.markdown(
                f"""
<div class="eco-kpi-card">
  <div class="eco-kpi-grid">
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">CAPA totales (vista filtrada + búsqueda)</div>
      <div class="eco-kpi-value">{total_capa}</div>
    </div>
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">CAPA abiertas</div>
      <div class="eco-kpi-value">{abiertas}</div>
      <div class="eco-kpi-chip">
        {_pct(abiertas, total_capa)}% del total
      </div>
    </div>
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">CAPA cerradas</div>
      <div class="eco-kpi-value">{cerradas}</div>
      <div class="eco-kpi-chip">
        {_pct(cerradas, total_capa)}% del total
      </div>
    </div>
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">Registros 8D vinculados (vista)</div>
      <div class="eco-kpi-value">{total_8d_vista}</div>
    </div>
  </div>
</div>
                """,
                unsafe_allow_html=True,
            )

            def lectura_indic_glob() -> str:
                backlog_pct = _pct(abiertas, total_capa) if total_capa > 0 else 0.0
                textos = []
                textos.append(
                    f"En la vista actual se gestionan **{total_capa} CAPA**, "
                    f"de las cuales **{abiertas} ({backlog_pct}% )** siguen abiertas "
                    f"y **{cerradas}** figuran como cerradas."
                )
                if total_8d_vista > 0:
                    textos.append(
                        f"Además, se observan **{total_8d_vista} informes 8D** vinculados, "
                        "lo que indica un nivel de documentación más profundo en parte "
                        "de las incidencias."
                    )
                textos.append(
                    "Este bloque resume la carga global del sistema y sirve como punto de partida "
                    "para las reuniones de seguimiento."
                )
                return " ".join(textos)

            st.markdown(
                f"<div class='eco-audit'>{lectura_indic_glob()}</div>",
                unsafe_allow_html=True,
            )

        # ------------------------------------------------------
        # 2) INDICADORES DE SLA
        # ------------------------------------------------------
        with t_sla:
            st.markdown("### ⏱️ Indicadores de SLA")

            sla_total_eval = sla_ok + sla_ko
            sla_txt_total = (
                f"{sla_total_eval} CAPA evaluadas"
                if sla_total_eval > 0
                else "Sin datos suficientes"
            )

            st.markdown(
                f"""
<div class="eco-kpi-card">
  <div class="eco-kpi-grid">
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">Cumplimiento global de SLA</div>
      <div class="eco-kpi-value">{sla_cumpl:.1f}%</div>
      <div class="eco-kpi-chip">
        Objetivo: {sla_target:.0f} {sla_mode}
      </div>
    </div>
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">CAPA dentro del SLA</div>
      <div class="eco-kpi-value">{sla_ok}</div>
      <div class="eco-kpi-chip">
        En plazo
      </div>
    </div>
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">CAPA fuera del SLA</div>
      <div class="eco-kpi-value">{sla_ko}</div>
      <div class="eco-kpi-chip">
        Fuera de plazo
      </div>
    </div>
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">Base de cálculo</div>
      <div class="eco-kpi-value">{sla_txt_total}</div>
    </div>
  </div>
</div>
                """,
                unsafe_allow_html=True,
            )

            def lectura_indic_sla() -> str:
                if sla_total_eval == 0:
                    return (
                        "Todavía no hay suficientes registros con fechas de apertura y cierre "
                        "para calcular un SLA robusto. Es importante completar estos campos "
                        "para poder evaluar la capacidad de respuesta del sistema."
                    )

                textos = []
                textos.append(
                    f"El **cumplimiento global de SLA** alcanza **{sla_cumpl:.1f}%** sobre "
                    f"**{sla_total_eval} CAPA** evaluadas. "
                    f"De ellas, **{sla_ok}** se cerraron dentro del objetivo de "
                    f"**{sla_target:.0f} {sla_mode}**, mientras que **{sla_ko}** "
                    "superaron dicho umbral."
                )

                if sla_cumpl >= 90:
                    textos.append(
                        "Este nivel refleja una **muy buena capacidad de respuesta**, "
                        "propia de un sistema de gestión maduro."
                    )
                elif sla_cumpl >= 70:
                    textos.append(
                        "El cumplimiento es **aceptable**, pero muestra margen de mejora, "
                        "especialmente en las CAPA que exceden el plazo objetivo."
                    )
                else:
                    textos.append(
                        "El nivel de cumplimiento es **bajo**, lo que indica que una parte "
                        "importante de las incidencias no se resuelven en el tiempo objetivo."
                    )

                textos.append(
                    "Técnicamente, este indicador se calcula como el porcentaje de CAPA cerradas "
                    "dentro del tiempo objetivo sobre el total de CAPA con fechas válidas."
                )
                return " ".join(textos)

            st.markdown(
                f"<div class='eco-audit'>{lectura_indic_sla()}</div>",
                unsafe_allow_html=True,
            )

        # ------------------------------------------------------
        # 3) INDICADORES DE REINCIDENCIAS
        # ------------------------------------------------------
        with t_reinc:
            st.markdown("### 🔁 Indicadores de reincidencias")

            st.markdown(
                f"""
<div class="eco-kpi-card">
  <div class="eco-kpi-grid">
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">Reincidencias 30 días (RC30)</div>
      <div class="eco-kpi-value">{r30_global}</div>
    </div>
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">Reincidencias 60 días (RC60)</div>
      <div class="eco-kpi-value">{r60_global}</div>
    </div>
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">CAPA con acciones correctivas (base de análisis)</div>
      <div class="eco-kpi-value">{acciones_corr}</div>
    </div>
  </div>
</div>
                """,
                unsafe_allow_html=True,
            )

            def lectura_indic_reinc() -> str:
                textos = []
                textos.append(
                    f"En la vista actual se identifican **{r30_global} reincidencias a 30 días** "
                    f"y **{r60_global} a 60 días**. Estas métricas permiten evaluar si los "
                    "problemas tienden a reaparecer tras aplicar acciones correctivas."
                )
                if acciones_corr > 0:
                    textos.append(
                        f"El volumen de reincidencias se interpreta siempre en relación con las "
                        f"**{acciones_corr} CAPA** que tienen acción correctiva registrada."
                    )
                else:
                    textos.append(
                        "Actualmente no hay acciones correctivas registradas; por ello, las "
                        "reincidencias deben interpretarse con cautela."
                    )
                textos.append(
                    "Técnicamente, RC30 y RC60 se calculan agrupando incidencias por su clave "
                    "de reincidencia y midiendo cuántas reaparecen en esas ventanas temporales."
                )
                return " ".join(textos)

            st.markdown(
                f"<div class='eco-audit'>{lectura_indic_reinc()}</div>",
                unsafe_allow_html=True,
            )

        # ------------------------------------------------------
        # 4) INDICADORES DE ACCIONES
        # ------------------------------------------------------
        with t_acc:
            st.markdown("### 🧩 Indicadores de acciones (contención / correctiva)")

            st.markdown(
                f"""
<div class="eco-kpi-card">
  <div class="eco-kpi-grid">
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">Acciones de contención (n≠0)</div>
      <div class="eco-kpi-value">{acciones_cont}</div>
      <div class="eco-kpi-chip">
        {_pct(acciones_cont, total_capa)}% de CAPA
      </div>
    </div>
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">Acciones correctivas (n≠0)</div>
      <div class="eco-kpi-value">{acciones_corr}</div>
      <div class="eco-kpi-chip">
        {_pct(acciones_corr, total_capa)}% de CAPA
      </div>
    </div>
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">Total de acciones registradas</div>
      <div class="eco-kpi-value">{acciones_totales}</div>
    </div>
  </div>
</div>
                """,
                unsafe_allow_html=True,
            )

            def lectura_indic_acc() -> str:
                textos = []
                textos.append(
                    f"De las **{total_capa} CAPA** analizadas, **{acciones_cont}** cuentan con "
                    "alguna **acción de contención** y "
                    f"**{acciones_corr}** disponen de al menos una **acción correctiva**."
                )
                textos.append(
                    "Una baja cobertura de acciones puede indicar debilidad en el tratamiento "
                    "formal de los problemas."
                )
                textos.append(
                    "Estos indicadores se obtienen revisando si los campos de acciones "
                    "contienen información no vacía para cada CAPA."
                )
                return " ".join(textos)

            st.markdown(
                f"<div class='eco-audit'>{lectura_indic_acc()}</div>",
                unsafe_allow_html=True,
            )

        # ------------------------------------------------------
        # 5) INDICADORES DE RIESGO & BACKLOG
        # ------------------------------------------------------
        with t_risk:
            st.markdown("### 🚦 Riesgo global y backlog")

            backlog_pct = _pct(abiertas, total_capa) if total_capa > 0 else 0.0

            st.markdown(
                f"""
<div class="eco-kpi-card">
  <div class="eco-kpi-grid">
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">Indicador global de riesgo</div>
      <div class="eco-kpi-value">{risk_global_score:.1f} / 100</div>
      <div class="eco-kpi-chip">
        Nivel: {risk_global_level}
      </div>
    </div>
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">% CAPA abiertas (backlog)</div>
      <div class="eco-kpi-value">{backlog_pct:.1f}%</div>
    </div>
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">Reincidencias 60 días (componente de riesgo)</div>
      <div class="eco-kpi-value">{r60_global}</div>
    </div>
  </div>
</div>
                """,
                unsafe_allow_html=True,
            )

            def lectura_indic_risk() -> str:
                textos = []
                textos.append(
                    f"El **indicador global de riesgo** se sitúa en **{risk_global_score:.1f} / 100**, "
                    f"clasificado como **{risk_global_level}**."
                )
                textos.append(
                    f"El cálculo pondera principalmente el **% de CAPA abiertas** "
                    f"({backlog_pct:.1f}%), el **incumplimiento de SLA** y las "
                    f"**reincidencias a 60 días ({r60_global})**."
                )
                textos.append(
                    "Este índice resume el nivel de tensión del sistema y ayuda a decidir "
                    "la urgencia de las acciones a nivel gerencial."
                )
                return " ".join(textos)

            st.markdown(
                f"<div class='eco-audit'>{lectura_indic_risk()}</div>",
                unsafe_allow_html=True,
            )

        # ------------------------------------------------------
        # 6) INDICADORES DE COSTOS
        # ------------------------------------------------------
        with t_cost:
            st.markdown("### 💶 Indicadores de costos")

            coste_medio = cost_total / total_capa if total_capa > 0 else 0.0

            st.markdown(
                f"""
<div class="eco-kpi-card">
  <div class="eco-kpi-grid">
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">Coste total estimado</div>
      <div class="eco-kpi-value">
        {cost_total:,.2f} €
      </div>
    </div>
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">Coste medio por CAPA</div>
      <div class="eco-kpi-value">
        {coste_medio:,.2f} €
      </div>
    </div>
  </div>
</div>
                """,
                unsafe_allow_html=True,
            )

            def lectura_indic_cost() -> str:
                if cost_total <= 0:
                    return (
                        "Actualmente no se dispone de información suficiente en los campos de "
                        "costes para calcular indicadores económicos robustos. Es recomendable "
                        "reforzar el registro de esta información."
                    )
                textos = []
                textos.append(
                    f"El **coste total estimado** asociado a las incidencias de la vista actual "
                    f"es de **{cost_total:,.2f} €**, con un **coste medio por CAPA** de "
                    f"aproximadamente **{coste_medio:,.2f} €**."
                )
                textos.append(
                    "Estos indicadores permiten vincular el desempeño de la calidad con el "
                    "impacto económico directo y priorizar procesos o clientes de mayor coste."
                )
                return " ".join(textos)

            st.markdown(
                f"<div class='eco-audit'>{lectura_indic_cost()}</div>",
                unsafe_allow_html=True,
            )
# ──────────────────────────────────────────────────────────────
# TAB 5 – Costos de incidencias
# ──────────────────────────────────────────────────────────────
with tab_costos:
    st.markdown("#### 💶 Costos de las incidencias")

    # Usamos la vista ya filtrada + buscador global si existe
    df_cost = df_view_search if "df_view_search" in globals() else df_view
    if df_cost is None:
        df_cost = pd.DataFrame()

    if df_cost.empty:
        st.info(
            "Por ahora no hay registros en la vista actual. "
            "Cuando existan incidencias con información de costes, aquí verás "
            "un resumen económico detallado y su interpretación."
        )
    else:
        # ──────────────────────────────────────────────────────
        # 1) Cálculo robusto de costes (reposiciones + devoluciones)
        # ──────────────────────────────────────────────────────
        def _to_num_safe(series_or_none):
            if series_or_none is None:
                return pd.Series(0.0, index=df_cost.index)
            try:
                return (
                    series_or_none.astype(str)
                    .str.replace(".", "", regex=False)
                    .str.replace(",", ".", regex=False)
                    .astype(float)
                )
            except Exception:
                return pd.Series(0.0, index=df_cost.index)

        serie_rep = (
            _to_num_safe(df_cost[col_cost_rep])
            if col_cost_rep and col_cost_rep in df_cost.columns
            else pd.Series(0.0, index=df_cost.index)
        )
        serie_dev = (
            _to_num_safe(df_cost[col_cost_dev])
            if col_cost_dev and col_cost_dev in df_cost.columns
            else pd.Series(0.0, index=df_cost.index)
        )

        df_cost_local = df_cost.copy()
        df_cost_local["_costo_reposicion"] = serie_rep
        df_cost_local["_costo_devolucion"] = serie_dev
        df_cost_local["_costo_total"] = (
            df_cost_local["_costo_reposicion"] + df_cost_local["_costo_devolucion"]
        )

        total_registros_cost = int(len(df_cost_local))
        coste_total_local = float(df_cost_local["_costo_total"].sum())
        coste_medio_local = (
            coste_total_local / total_registros_cost if total_registros_cost > 0 else 0.0
        )

        # ──────────────────────────────────────────────────────
        # 2) Tarjeta principal de KPIs económicos
        # ──────────────────────────────────────────────────────
        st.markdown(
            f"""
<div class="eco-kpi-card">
  <div class="eco-kpi-grid">
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">Incidencias (vista filtrada + búsqueda)</div>
      <div class="eco-kpi-value">{total_registros_cost}</div>
    </div>
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">Coste total estimado</div>
      <div class="eco-kpi-value">
        {coste_total_local:,.2f} €
      </div>
    </div>
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">Coste medio por incidencia</div>
      <div class="eco-kpi-value">
        {coste_medio_local:,.2f} €
      </div>
    </div>
    <div class="eco-kpi-item">
      <div class="eco-kpi-label">Peso aprox. en el total global del dashboard</div>
      <div class="eco-kpi-value">
        {(_pct(coste_total_local, cost_total) if cost_total > 0 else 0):.1f}%
      </div>
      <div class="eco-kpi-chip">
        Frente al coste total estimado (todas las vistas)
      </div>
    </div>
  </div>
</div>
            """,
            unsafe_allow_html=True,
        )

        # ──────────────────────────────────────────────────────
        # 3) Gráfico 1 · Coste total por tipo de coste
        # ──────────────────────────────────────────────────────
        st.markdown("##### 📊 Coste total por tipo de coste")

        df_cost_tipo = pd.DataFrame(
            {
                "Tipo de coste": ["Reposición", "Devolución"],
                "Coste total": [
                    float(df_cost_local["_costo_reposicion"].sum()),
                    float(df_cost_local["_costo_devolucion"].sum()),
                ],
            }
        )

        if df_cost_tipo["Coste total"].sum() == 0:
            st.info(
                "No se dispone de valores numéricos en los campos de reposición/devolución "
                "para construir un gráfico por tipo de coste."
            )
        else:
            try:
                fig_cost_tipo = px.bar(
                    df_cost_tipo,
                    x="Tipo de coste",
                    y="Coste total",
                    text="Coste total",
                    title="Coste total por tipo de coste (reposiciones vs devoluciones)",
                )
                fig_cost_tipo.update_traces(
                    texttemplate="%{text:,.2f} €",
                    textposition="outside",
                    cliponaxis=False,
                )
                fig_cost_tipo.update_layout(
                    margin=dict(l=10, r=10, t=50, b=40),
                    height=420,
                    yaxis_title="Coste total (€)",
                )
                st.plotly_chart(fig_cost_tipo, width="stretch")
                # Caption bajo el gráfico (app)
                st.caption(
                    "Figura 1. Distribución del coste total de las incidencias entre "
                    "reposiciones internas y devoluciones."
                )
                # Registro para el informe en Word
                if "register_figure" in globals():
                    register_figure(
                        fig=fig_cost_tipo,
                        fig_id="costos_tipo",
                        title="Coste total por tipo de coste (reposiciones vs devoluciones)",
                        caption=(
                            "Distribución del coste total de las incidencias entre "
                            "reposiciones internas y devoluciones."
                        ),
                        source_tab="Costos",
                    )
            except Exception as e:
                st.warning(f"No se pudo renderizar el gráfico por tipo de coste: {e}")

        # ──────────────────────────────────────────────────────
        # 4) Gráfico 2 · Coste total por tipo de incidencia/defecto (si existe columna)
        # ──────────────────────────────────────────────────────
        st.markdown("##### 📊 Coste total por tipo de incidencia / defecto")

        col_def_cost = _find_col_any(
            df_cost_local,
            [
                "defecto",
                "defectos",
                "tipo incidencia",
                "incidencia",
                "motivo",
                "descripción detallada de la incidencia",
                "descripcion detallada de la incidencia",
            ],
        )

        df_cost_cat = pd.DataFrame()  # inicializamos por seguridad

        if not col_def_cost:
            st.info(
                "No se encontró una columna clara de tipo de incidencia/defecto para "
                "analizar el coste por categoría. Revisa el diccionario de datos si "
                "quieres activar este análisis."
            )
        else:
            serie_def_cost = (
                df_cost_local[col_def_cost]
                .astype(str)
                .str.strip()
                .replace({"nan": "", "None": ""})
            )
            mask_valid_def = serie_def_cost != ""

            df_cost_cat = (
                df_cost_local.loc[mask_valid_def, ["_costo_total"]]
                .assign(Defecto=serie_def_cost[mask_valid_def])
                .groupby("Defecto", as_index=False)["_costo_total"]
                .sum()
            )

            df_cost_cat = df_cost_cat.sort_values("_costo_total", ascending=False)

            if df_cost_cat.empty or df_cost_cat["_costo_total"].sum() == 0:
                st.info(
                    "La columna detectada como defecto/tipo de incidencia no contiene "
                    "valores significativos asociados a costes en la vista actual."
                )
            else:
                try:
                    fig_cost_cat = px.bar(
                        df_cost_cat,
                        x="Defecto",
                        y="_costo_total",
                        text="_costo_total",
                        title="Coste total por tipo de incidencia / defecto",
                    )
                    fig_cost_cat.update_traces(
                        texttemplate="%{text:,.2f} €",
                        textposition="outside",
                        cliponaxis=False,
                    )
                    fig_cost_cat.update_layout(
                        margin=dict(l=10, r=10, t=60, b=160),
                        height=450,
                        xaxis_title=None,
                        yaxis_title="Coste total (€)",
                    )
                    fig_cost_cat.update_xaxes(tickangle=-30, automargin=True)
                    st.plotly_chart(fig_cost_cat, width="stretch")
                    # Caption bajo el gráfico (app)
                    st.caption(
                        "Figura 2. Coste total de las incidencias agrupado por tipo "
                        "de incidencia o defecto principal."
                    )
                    # Registro para el informe en Word
                    if "register_figure" in globals():
                        register_figure(
                            fig=fig_cost_cat,
                            fig_id="costos_por_defecto",
                            title="Coste total por tipo de incidencia / defecto",
                            caption=(
                                "Coste total de las incidencias agrupado por tipo de incidencia "
                                "o defecto principal."
                            ),
                            source_tab="Costos",
                        )
                except Exception as e:
                    st.warning(f"No se pudo renderizar el gráfico de costes por defecto: {e}")

        # ──────────────────────────────────────────────────────
        # 5) Lectura automática de costes (estilo auditoría)
        # ──────────────────────────────────────────────────────
        def lectura_costos_detallada() -> str:
            if total_registros_cost == 0:
                return (
                    "No hay incidencias en la vista actual con información de costes. "
                    "Para explotar este módulo es necesario registrar los importes de "
                    "reposiciones y/o devoluciones asociadas a cada CAPA."
                )

            textos = []
            textos.append(
                f"En la vista actual se analizan **{total_registros_cost} incidencias** "
                f"con información económica, que generan un **coste total estimado de "
                f"{coste_total_local:,.2f} €**. El **coste medio por incidencia** se sitúa "
                f"en torno a **{coste_medio_local:,.2f} €**."
            )

            # Detalle de reposición vs devolución
            total_rep = float(df_cost_local["_costo_reposicion"].sum())
            total_dev = float(df_cost_local["_costo_devolucion"].sum())
            total_all = total_rep + total_dev

            if total_all > 0:
                p_rep = _pct(total_rep, total_all)
                p_dev = _pct(total_dev, total_all)
                textos.append(
                    f"Del coste total, aproximadamente **{p_rep:.1f}%** corresponde a "
                    f"**reposiciones** y **{p_dev:.1f}%** a **devoluciones**. "
                    "Esta descomposición permite identificar si el impacto económico proviene "
                    "más de la reposición interna de producto o de la relación con el cliente."
                )

            # Si hay desglose por defecto
            if col_def_cost and not df_cost_cat.empty:
                defecto_top = df_cost_cat.iloc[0]["Defecto"]
                coste_top = float(df_cost_cat.iloc[0]["_costo_total"])
                p_top = _pct(coste_top, coste_total_local) if coste_total_local > 0 else 0.0

                textos.append(
                    f"Al analizar el **coste por tipo de incidencia**, el defecto "
                    f"**{defecto_top}** concentra aproximadamente **{coste_top:,.2f} €**, "
                    f"equivalentes a **{p_top:.1f}%** del coste total. "
                    "Este tipo de análisis es clave para priorizar proyectos de mejora "
                    "en aquellos problemas que más dinero consumen."
                )

            textos.append(
                "Desde un punto de vista técnico, estos indicadores se construyen a partir "
                "de los campos de coste de reposición y devolución, transformándolos a valores "
                "numéricos y agregándolos por incidencia y, cuando es posible, por categoría "
                "de defecto. De esta forma, el módulo de costes conecta directamente la calidad "
                "con el impacto económico."
            )

            return " ".join(textos)

        with st.expander(
            "🧾 Lectura automática de costos (estilo auditoría)",
            expanded=False,
        ):
            st.markdown(
                f"<div class='eco-audit'>{lectura_costos_detallada()}</div>",
                unsafe_allow_html=True,
            )
# ──────────────────────────────────────────────────────────────
# TAB 7 – Tendencias & pronósticos
# ──────────────────────────────────────────────────────────────
with tab_tendencias:
    st.markdown("#### 📈 Tendencias & pronóstico de incidencias y costos")

    # Usamos la vista ya filtrada + buscador global si existe
    df_trend = df_view_search if "df_view_search" in globals() else df_view
    if df_trend is None:
        df_trend = pd.DataFrame()

    if df_trend.empty:
        st.info(
            "Por ahora no hay registros en la vista actual. "
            "Cuando existan incidencias, esta pestaña mostrará la evolución temporal "
            "de CAPA, los costes asociados y un pequeño pronóstico automático."
        )
    else:
        # 0) Detectamos columna de fecha para trabajar por mes
        col_fecha = _find_col_any(
            df_trend,
            [
                "fecha creacion",
                "fecha creación",
                "fecha detección",
                "fecha de detección",
                "fecha",
                "fecha apertura",
                "fecha de apertura",
            ],
        )

        if not col_fecha:
            st.info(
                "No se ha detectado una columna de fecha clara para construir las tendencias. "
                "Revisa el diccionario de datos y asegúrate de incluir una fecha de creación/detección."
            )
        else:
            df_trend = df_trend.copy()
            df_trend["_fecha"] = pd.to_datetime(df_trend[col_fecha], errors="coerce")
            df_trend = df_trend.dropna(subset=["_fecha"])

            if df_trend.empty:
                st.info(
                    "No se encontraron fechas válidas para construir tendencias. "
                    "Revisa el formato de las fechas en la base de datos."
                )
            else:
                # Agrupamos por año-mes
                df_trend["_año_mes"] = df_trend["_fecha"].dt.to_period("M")
                df_trend_mes = (
                    df_trend.groupby("_año_mes")
                    .size()
                    .reset_index(name="CAPA_mes")
                    .sort_values("_año_mes")
                )
                df_trend_mes["mes"] = df_trend_mes["_año_mes"].astype(str)

                # ────────────────────────────────────────────────────
                # 1) Tendencia mensual del volumen de CAPA
                # ────────────────────────────────────────────────────
                st.markdown("### 📌 Evolución mensual de CAPA registradas")

                if df_trend_mes.empty:
                    st.info("No hay suficientes datos temporales para construir la serie mensual.")
                else:
                    fig_capa = px.line(
                        df_trend_mes,
                        x="mes",
                        y="CAPA_mes",
                        markers=True,
                        title="CAPA registradas por mes",
                    )
                    fig_capa.update_layout(
                        xaxis_title="Mes",
                        yaxis_title="Nº de CAPA registradas",
                        height=430,
                        margin=dict(l=10, r=10, t=40, b=40),
                    )
                    st.plotly_chart(fig_capa, width='stretch')

                    # Caption en la app
                    st.caption(
                        "Figura 1. Evolución mensual del número de incidencias CAPA registradas "
                        "en el sistema."
                    )

                    # Registro para el informe Word
                    if "register_figure" in globals():
                        register_figure(
                            fig=fig_capa,
                            fig_id="tend_capa_mensual",
                            title="CAPA registradas por mes",
                            caption=(
                                "Evolución mensual del número de incidencias CAPA registradas "
                                "en el sistema."
                            ),
                            source_tab="Tendencias & pronósticos",
                        )

                    def lectura_tend_capa() -> str:
                        textos = []
                        total_periodos = df_trend_mes.shape[0]
                        total_capa_hist = int(df_trend_mes["CAPA_mes"].sum())

                        textos.append(
                            f"En el periodo analizado se registran **{total_capa_hist} CAPA** "
                            f"distribuidas en **{total_periodos} meses**."
                        )

                        if total_periodos >= 2:
                            primera = int(df_trend_mes["CAPA_mes"].iloc[0])
                            ultima = int(df_trend_mes["CAPA_mes"].iloc[-1])
                            delta = ultima - primera
                            delta_pct = _pct(delta, primera) if primera > 0 else 0.0
                            textos.append(
                                f"El volumen mensual pasa de **{primera} CAPA** en el primer mes "
                                f"a **{ultima} CAPA** en el último, lo que supone una variación "
                                f"de **{delta:+d} casos ({delta_pct:+.1f}%)**."
                            )
                        else:
                            textos.append(
                                "Solo se dispone de un mes de información, por lo que aún no es "
                                "posible evaluar una tendencia clara."
                            )

                        textos.append(
                            "Este gráfico permite identificar meses pico de incidencias, estacionalidad "
                            "y posibles efectos de cambios de proceso o lanzamientos, y sirve de base "
                            "para investigar más en detalle por cliente, producto o área."
                        )
                        return " ".join(textos)

                    with st.expander(
                        "🧾 Lectura automática de la tendencia de CAPA (estilo auditoría)",
                        expanded=False,
                    ):
                        st.markdown(
                            f"<div class='eco-audit'>{lectura_tend_capa()}</div>",
                            unsafe_allow_html=True,
                        )

                # ────────────────────────────────────────────────────
                # 2) Tendencia mensual de costes (si hay datos)
                # ────────────────────────────────────────────────────
                st.markdown("### 💶 Tendencia mensual de costes asociados (si aplica)")

                # Solo calculamos costes si existen columnas definidas globalmente
                df_cost = df_trend.copy()
                tiene_costes = False

                def _to_num_cost(s) -> float:
                    try:
                        return float(str(s).replace(".", "").replace(",", "."))
                    except Exception:
                        return 0.0

                if col_cost_rep and col_cost_rep in df_cost.columns:
                    df_cost["_cost_rep"] = df_cost[col_cost_rep].map(_to_num_cost)
                    tiene_costes = True
                else:
                    df_cost["_cost_rep"] = 0.0

                if col_cost_dev and col_cost_dev in df_cost.columns:
                    df_cost["_cost_dev"] = df_cost[col_cost_dev].map(_to_num_cost)
                    tiene_costes = True
                else:
                    df_cost["_cost_dev"] = 0.0

                df_cost["_coste_total"] = df_cost["_cost_rep"] + df_cost["_cost_dev"]

                df_cost_mes = (
                    df_cost.groupby("_año_mes")["_coste_total"]
                    .sum()
                    .reset_index(name="Coste_total_mes")
                    .sort_values("_año_mes")
                )
                df_cost_mes["mes"] = df_cost_mes["_año_mes"].astype(str)

                if (not tiene_costes) or df_cost_mes["Coste_total_mes"].sum() == 0:
                    st.info(
                        "No se dispone de datos significativos en las columnas de coste para "
                        "construir una tendencia económica. Revisa el registro de costes de "
                        "reposiciones/devoluciones."
                    )
                else:
                    fig_cost = px.bar(
                        df_cost_mes,
                        x="mes",
                        y="Coste_total_mes",
                        title="Coste total estimado por mes",
                    )
                    fig_cost.update_layout(
                        xaxis_title="Mes",
                        yaxis_title="Coste total (unidades monetarias)",
                        height=430,
                        margin=dict(l=10, r=10, t=40, b=40),
                    )
                    st.plotly_chart(fig_cost, width="stretch")

                    # Caption en la app
                    st.caption(
                        "Figura 2. Evolución mensual del coste total estimado asociado a las "
                        "incidencias registradas."
                    )

                    # Registro para el informe Word
                    if "register_figure" in globals():
                        register_figure(
                            fig=fig_cost,
                            fig_id="tend_coste_mensual",
                            title="Coste total estimado por mes",
                            caption=(
                                "Evolución mensual del coste total estimado asociado a las "
                                "incidencias registradas."
                            ),
                            source_tab="Tendencias & pronósticos",
                        )

                    def lectura_tend_costes() -> str:
                        textos = []
                        total_cost_hist = float(df_cost_mes["Coste_total_mes"].sum())
                        textos.append(
                            f"En el periodo analizado, el **coste total estimado** asociado a las incidencias "
                            f"asciende aproximadamente a **{total_cost_hist:,.2f}** unidades monetarias."
                        )
                        if df_cost_mes.shape[0] >= 2:
                            c_first = float(df_cost_mes["Coste_total_mes"].iloc[0])
                            c_last = float(df_cost_mes["Coste_total_mes"].iloc[-1])
                            c_delta = c_last - c_first
                            c_delta_pct = _pct(c_delta, c_first) if c_first > 0 else 0.0
                            textos.append(
                                f"El coste mensual pasa de **{c_first:,.2f}** en el primer mes "
                                f"a **{c_last:,.2f}** en el último, lo que supone una variación "
                                f"de **{c_delta:+.2f} ({c_delta_pct:+.1f}%)**."
                            )
                        else:
                            textos.append(
                                "Solo se dispone de un mes con costes registrados; todavía no es "
                                "posible evaluar una tendencia de reducción o incremento."
                            )

                        textos.append(
                            "Este gráfico permite vincular la evolución de las incidencias con su impacto "
                            "económico, identificando meses con picos de coste que conviene analizar en más "
                            "detalle (por cliente, producto, proceso, etc.)."
                        )
                        return " ".join(textos)

                    with st.expander(
                        "🧾 Lectura automática de la tendencia de costes (estilo auditoría)",
                        expanded=False,
                    ):
                        st.markdown(
                            f"<div class='eco-audit'>{lectura_tend_costes()}</div>",
                            unsafe_allow_html=True,
                        )

                # ────────────────────────────────────────────────────
                # 3) Pronóstico simple para el próximo periodo
                # ────────────────────────────────────────────────────
                st.markdown("### 🔮 Pronóstico simple del volumen de CAPA")

                if df_trend_mes.shape[0] < 3:
                    st.info(
                        "Se dispone de menos de 3 meses de datos; por ahora solo se muestra la tendencia histórica "
                        "sin pronóstico numérico. A medida que se acumulen meses, se habilitará un pronóstico simple."
                    )
                else:
                    # Usamos los últimos 3 meses para una regresión lineal muy simple
                    df_fore = df_trend_mes.copy()
                    ultimos = df_fore.tail(3)
                    y = ultimos["CAPA_mes"].to_numpy()
                    x = np.arange(len(y))

                    try:
                        coef = np.polyfit(x, y, 1)
                        next_x = len(y)
                        y_next = max(0, coef[0] * next_x + coef[1])

                        ultimo_periodo = df_fore["_año_mes"].max()
                        next_periodo = ultimo_periodo + 1

                        next_row = pd.DataFrame(
                            {
                                "_año_mes": [next_periodo],
                                "CAPA_mes": [y_next],
                                "mes": [str(next_periodo)],
                            }
                        )
                        df_fore["tipo"] = "Histórico"
                        next_row["tipo"] = "Pronóstico"

                        df_forecast = pd.concat([df_fore, next_row], ignore_index=True)

                        fig_fore = px.line(
                            df_forecast,
                            x="mes",
                            y="CAPA_mes",
                            color="tipo",
                            markers=True,
                            title="Pronóstico simple del volumen de CAPA (próximo mes)",
                        )
                        fig_fore.update_layout(
                            xaxis_title="Mes",
                            yaxis_title="Nº de CAPA",
                            height=430,
                            margin=dict(l=10, r=10, t=40, b=40),
                        )
                        st.plotly_chart(fig_fore, width="stretch")

                        # Caption en la app
                        st.caption(
                            "Figura 3. Pronóstico lineal sencillo del volumen de CAPA para el "
                            "próximo mes, a partir de los últimos tres meses históricos."
                        )

                        # Registro para el informe Word
                        if "register_figure" in globals():
                            register_figure(
                                fig=fig_fore,
                                fig_id="pronostico_capa",
                                title="Pronóstico simple del volumen de CAPA (próximo mes)",
                                caption=(
                                    "Pronóstico lineal sencillo del volumen de CAPA para el próximo "
                                    "mes, calculado a partir de los últimos tres meses históricos."
                                ),
                                source_tab="Tendencias & pronósticos",
                            )

                        def lectura_pronostico() -> str:
                            textos = []
                            textos.append(
                                "Se ha aplicado un **modelo lineal muy simple** sobre los últimos 3 meses "
                                "para estimar el volumen de CAPA del próximo periodo."
                            )
                            textos.append(
                                f"El resultado sugiere un valor esperado en torno a **{y_next:.1f} CAPA** "
                                "para el próximo mes, que debe interpretarse como una referencia orientativa "
                                "y no como un compromiso de carga real."
                            )
                            textos.append(
                                "Este tipo de pronóstico es útil para anticipar necesidades de recursos, "
                                "pero debe contrastarse siempre con la planificación operativa, campañas, "
                                "cambios de proceso o eventos excepcionales."
                            )
                            return " ".join(textos)

                        with st.expander(
                            "🧾 Lectura automática del pronóstico (estilo auditoría)",
                            expanded=False,
                        ):
                            st.markdown(
                                f"<div class='eco-audit'>{lectura_pronostico()}</div>",
                                unsafe_allow_html=True,
                            )
                    except Exception:
                        st.info(
                            "No fue posible calcular el pronóstico simple con los datos actuales. "
                            "Se mantiene únicamente la visualización histórica."
                        )
# ─────────────────────────────────────────────
# PESTAÑA: PLAN DE ACCIONES Y MEJORA CONTINUA
# ─────────────────────────────────────────────
with tab_plan_ia:  # ← ajusta este nombre al de tu pestaña real
    st.markdown("## 🧭 Plan de acciones y mejora continua")

    # Usamos la BDCAPA que ya está cargada desde SQLite
    df_base = df_capa.copy() if "df_capa" in locals() else pd.DataFrame()

    if df_base.empty:
        st.info("No hay datos en BDCAPA para generar el plan de acciones.")
    else:
        # ─────────────────────────────────────
        # 1. Detectar columnas clave de BDCAPA
        # ─────────────────────────────────────
        col_capa = _find_col_any(df_base, ["nº capa", "no capa", "numero capa", "n capa", "capa"])
        col_estatus = _find_col_any(df_base, ["estatus", "status", "estado"])
        col_cliente = _find_col_any(df_base, ["cliente"])
        col_pedido = _find_col_any(df_base, ["pedido origen", "pedido", "nombre del pedido"])
        col_inc = _find_col_any(df_base, ["incidencia", "descripcion", "descripción"])
        col_causa = _find_col_any(
            df_base,
            [
                "causa raiz",
                "causa raíz",
                "descripcion resumida causa raiz",
                "descripcion resumida causa raiz- (extendida en informe 8d)",
            ],
        )
        col_acc_cont = _find_col_any(df_base, ["accion contencion", "acción contención"])
        col_acc_corr = _find_col_any(df_base, ["accion correccion", "acción corrección"])
        col_repo = _find_col_any(df_base, ["n reposicion", "nº reposición", "numero de reposicion"])
        col_e_repo = _find_col_any(df_base, ["€ reposición", "e reposicion", "costo reposicion"])
        col_e_dev = _find_col_any(df_base, ["€ devolucion", "e devolucion", "costo devolucion"])
        col_fecha_crea = _find_col_any(df_base, ["fecha creacion", "fecha creación", "fecha capa"])
        col_resp = _find_col_any(df_base, ["responsable"])
        col_tecnico = _find_col_any(df_base, ["tecnico", "técnico"])
        col_enlace_8d = _find_col_any(df_base, ["enlace informe 8d"])

        # ─────────────────────────────────────
        # 2. Controles superiores / Buscador
        # ─────────────────────────────────────
        st.markdown("### 🔍 Filtros y buscador de incidencias")

        with st.container():
            c1, c2, c3 = st.columns([2, 2, 2])

            with c1:
                texto_buscar = st.text_input(
                    "Buscar por Nº CAPA, pedido, cliente, incidencia, causa raíz…",
                    value="",
                    placeholder="Ej: CP-12, BOMBA, FUGA, cliente X…",
                    key="buscar_plan_acciones",
                )

            with c2:
                if col_cliente and not df_base.empty:
                    opciones_cliente = (
                        df_base[col_cliente]
                        .dropna()
                        .astype(str)
                        .sort_values()
                        .unique()
                        .tolist()
                    )
                else:
                    opciones_cliente = []
                clientes_sel = st.multiselect(
                    "Filtrar por cliente",
                    opciones_cliente,
                    default=[],
                )

            with c3:
                if col_estatus and not df_base.empty:
                    opciones_est = (
                        df_base[col_estatus]
                        .dropna()
                        .astype(str)
                        .sort_values()
                        .unique()
                        .tolist()
                    )
                else:
                    opciones_est = []
                estatus_sel = st.multiselect(
                    "Filtrar por estatus",
                    opciones_est,
                    default=[],
                )

        df_actions = df_base.copy()

        # Filtro por cliente
        if clientes_sel and col_cliente:
            df_actions = df_actions[df_actions[col_cliente].astype(str).isin(clientes_sel)]

        # Filtro por estatus
        if estatus_sel and col_estatus:
            df_actions = df_actions[df_actions[col_estatus].astype(str).isin(estatus_sel)]

        # Buscador de texto global
        if texto_buscar.strip():
            s = texto_buscar.strip().lower()
            mask = pd.Series(False, index=df_actions.index)
            for c in df_actions.columns:
                mask = mask | df_actions[c].astype(str).str.lower().str.contains(s, na=False)
            df_actions = df_actions[mask]

        if df_actions.empty:
            st.warning("No se encontraron incidencias que coincidan con los filtros aplicados.")
        else:
            # ─────────────────────────────────────
            # 3. Columnas inteligentes
            # ─────────────────────────────────────
            from datetime import datetime

            def _safe_num(x):
                try:
                    return float(str(x).replace(",", "."))
                except Exception:
                    return 0.0

            def _prioridad_automatica(row) -> str:
                monto = 0.0
                if col_e_repo and pd.notna(row.get(col_e_repo, None)):
                    monto += _safe_num(row[col_e_repo])
                if col_e_dev and pd.notna(row.get(col_e_dev, None)):
                    monto += _safe_num(row[col_e_dev])

                texto = ""
                if col_inc and pd.notna(row.get(col_inc, None)):
                    texto += str(row[col_inc]) + " "
                if col_causa and pd.notna(row.get(col_causa, None)):
                    texto += str(row[col_causa])
                t = texto.lower()

                if "seguridad" in t or "accidente" in t or monto >= 2000:
                    return "Alta"
                if (
                    "reclamacion" in t
                    or "reclamación" in t
                    or "devolucion" in t
                    or "devolución" in t
                    or monto >= 800
                ):
                    return "Media"
                return "Baja"

            def _riesgo_automatico(row) -> str:
                monto = 0.0
                if col_e_repo and pd.notna(row.get(col_e_repo, None)):
                    monto += _safe_num(row[col_e_repo])
                if col_e_dev and pd.notna(row.get(col_e_dev, None)):
                    monto += _safe_num(row[col_e_dev])

                t = ""
                if col_inc and pd.notna(row.get(col_inc, None)):
                    t += str(row[col_inc]).lower() + " "
                if col_causa and pd.notna(row.get(col_causa, None)):
                    t += str(row[col_causa]).lower()

                riesgo = "Medio"
                if "lesion" in t or "lesión" in t or "seguridad" in t or "accidente" in t:
                    riesgo = "Muy alto"
                elif "parada" in t or "paralizacion" in t or "paralización" in t or monto > 3000:
                    riesgo = "Muy alto"
                elif "retrabajo" in t or "devolucion" in t or "devolución" in t or monto > 1000:
                    riesgo = "Alto"
                elif monto < 300:
                    riesgo = "Bajo"

                return riesgo

            def _riesgo_score(prioridad: str, riesgo: str) -> int:
                p_map = {"Alta": 3, "Media": 2, "Baja": 1}
                r_map = {"Muy alto": 4, "Alto": 3, "Medio": 2, "Bajo": 1}
                return p_map.get(prioridad, 1) * r_map.get(riesgo, 1)

            def _estado_calculado(row) -> str:
                if col_estatus and pd.notna(row.get(col_estatus, None)):
                    est = str(row[col_estatus]).strip().lower()
                    if "cerrad" in est or "complet" in est:
                        return "Completada"
                    if "curso" in est or "proceso" in est:
                        return "En curso"
                return "Pendiente"

            def _responsable_sugerido(row) -> str:
                t = ""
                if col_inc and pd.notna(row.get(col_inc, None)):
                    t += str(row[col_inc]).lower() + " "
                if col_causa and pd.notna(row.get(col_causa, None)):
                    t += str(row[col_causa]).lower()

                if "medida" in t or "dimensión" in t or "dimension" in t or "especificacion" in t:
                    return "Calidad"
                if "entrega" in t or "retraso" in t or "logistica" in t:
                    return "Logística"
                if (
                    "proceso" in t
                    or "linea" in t
                    or "línea" in t
                    or "maquina" in t
                    or "máquina" in t
                ):
                    return "Producción"
                if "cliente" in t or "reclamacion" in t or "reclamación" in t:
                    return "Servicio al cliente"
                return "Responsable a definir"

            def _accion_sugerida(row) -> str:
                capa = row[col_capa] if col_capa else ""
                inc = row[col_inc] if col_inc else ""
                causa = row[col_causa] if col_causa else ""
                cliente = row[col_cliente] if col_cliente else ""
                prioridad = row.get("Prioridad (automática)", "")
                riesgo = row.get("Riesgo (automático)", "")
                resp_sug = row.get("Responsable sugerido", "")

                partes = []
                if capa:
                    partes.append(f"Para la CAPA {capa}")
                else:
                    partes.append("Para esta incidencia")

                if cliente:
                    partes.append(f"del cliente {cliente}")

                base = " ".join(partes).strip()
                cuerpo = "se recomienda:"
                sugerencia = []

                if "seguridad" in str(inc).lower() or "accidente" in str(inc).lower():
                    sugerencia.append(
                        "detener temporalmente el proceso afectado y evaluar riesgos de seguridad."
                    )
                elif "medida" in str(causa).lower() or "especificacion" in str(causa).lower():
                    sugerencia.append(
                        "revisar especificaciones técnicas, calibrar equipos de medición y validar primeras piezas."
                    )
                elif "entrega" in str(inc).lower() or "retraso" in str(inc).lower():
                    sugerencia.append(
                        "analizar cuellos de botella en logística y ajustar planificación o proveedores clave."
                    )
                else:
                    sugerencia.append(
                        "realizar una revisión rápida del proceso involucrado y definir una acción correctiva concreta con fecha y responsable."
                    )

                if prioridad == "Alta" or riesgo in ("Muy alto", "Alto"):
                    sugerencia.append(" Dar seguimiento en las próximas 24–48 horas.")
                else:
                    sugerencia.append(" Programar el seguimiento en la próxima reunión operativa.")

                if resp_sug and resp_sug != "Responsable a definir":
                    sugerencia.append(f" Responsable sugerido: {resp_sug}.")

                return f"{base}, {cuerpo} " + " ".join(sugerencia)

            # Añadir columnas inteligentes
            df_actions = df_actions.copy()
            df_actions["Estado (automático)"] = df_actions.apply(_estado_calculado, axis=1)
            df_actions["Prioridad (automática)"] = df_actions.apply(_prioridad_automatica, axis=1)
            df_actions["Riesgo (automático)"] = df_actions.apply(_riesgo_automatico, axis=1)
            df_actions["Score riesgo"] = df_actions.apply(
                lambda r: _riesgo_score(r["Prioridad (automática)"], r["Riesgo (automático)"]),
                axis=1,
            )
            df_actions["Responsable sugerido"] = df_actions.apply(_responsable_sugerido, axis=1)
            df_actions["Acción sugerida (texto IA simple)"] = df_actions.apply(
                _accion_sugerida, axis=1
            )

            # Días abiertos
            hoy = datetime.today().date()
            dias_abiertos = []
            for _, row in df_actions.iterrows():
                if col_fecha_crea and pd.notna(row.get(col_fecha_crea, None)):
                    try:
                        f = row[col_fecha_crea]
                        if isinstance(f, pd.Timestamp):
                            f = f.date()
                        elif isinstance(f, datetime):
                            f = f.date()
                        else:
                            f = pd.to_datetime(f).date()
                        dias_abiertos.append((hoy - f).days)
                    except Exception:
                        dias_abiertos.append(None)
                else:
                    dias_abiertos.append(None)
            df_actions["Días abiertos (aprox.)"] = dias_abiertos

            # ─────────────────────────────────────
            # 4. KPIs y alertas principales (con porcentajes)
            # ─────────────────────────────────────
            st.markdown("### 📊 Resumen ejecutivo del plan de acciones")

            total = len(df_actions)
            comp = (df_actions["Estado (automático)"] == "Completada").sum()
            pend = (df_actions["Estado (automático)"] != "Completada").sum()
            alto_riesgo = (
                df_actions["Riesgo (automático)"].isin(["Muy alto", "Alto"])
            ).sum()

            pct_comp = round(comp / total * 100, 1) if total else 0.0
            pct_pend = round(pend / total * 100, 1) if total else 0.0
            pct_alto_riesgo = round(alto_riesgo / total * 100, 1) if total else 0.0

            col_k1, col_k2, col_k3, col_k4 = st.columns(4)
            with col_k1:
                st.metric("Total CAPA en plan", total)
            with col_k2:
                st.metric("Acciones completadas", f"{comp} ({pct_comp}%)")
            with col_k3:
                st.metric("Acciones abiertas", f"{pend} ({pct_pend}%)")
            with col_k4:
                st.metric("CAPA de riesgo alto/muy alto", f"{alto_riesgo} ({pct_alto_riesgo}%)")

            # Lectura automática del resumen ejecutivo
            st.markdown("**📝 Lectura automática – Resumen ejecutivo**")
            st.write(
                f"Actualmente el plan de acciones contempla **{total}** CAPA. "
                f"De ellas, **{comp}** ({pct_comp}%) se encuentran **completadas**, "
                f"mientras que **{pend}** ({pct_pend}%) siguen **abiertas o en curso**. "
                f"Además, **{alto_riesgo}** CAPA ({pct_alto_riesgo}%) están clasificadas "
                f"como de **riesgo Alto o Muy alto**, lo que indica los casos que requieren "
                f"mayor foco en las próximas reuniones operativas."
            )

            # Alertas
            st.markdown("#### ⚠️ Alertas automáticas")
            alertas = []

            # CAPAs con riesgo muy alto
            df_riesgo_muy_alto = df_actions[df_actions["Riesgo (automático)"] == "Muy alto"]
            if not df_riesgo_muy_alto.empty:
                alertas.append(
                    f"- {len(df_riesgo_muy_alto)} CAPA clasificadas como **Riesgo Muy alto**."
                )

            # CAPAs muy antiguas (>30 días abiertas)
            df_abiertas = df_actions[df_actions["Estado (automático)"] != "Completada"]
            df_viejas = df_abiertas[df_abiertas["Días abiertos (aprox.)"].fillna(0) > 30]
            if not df_viejas.empty:
                alertas.append(
                    f"- {len(df_viejas)} CAPA abiertas desde hace más de **30 días**, "
                    "que conviene revisar y cerrar."
                )

            # CAPAs sin responsable
            if col_resp:
                df_sin_resp = df_actions[
                    df_actions[col_resp].isna()
                    | (df_actions[col_resp].astype(str).str.strip() == "")
                ]
                if not df_sin_resp.empty:
                    alertas.append(
                        f"- {len(df_sin_resp)} CAPA **sin responsable asignado**, "
                        "lo que puede retrasar la implementación de acciones."
                    )
            else:
                df_sin_resp = df_actions  # por si queremos usarlo luego

            if alertas:
                st.write("\n".join(alertas))
            else:
                st.success("No se encontraron alertas críticas según las reglas actuales.")

            # ─────────────────────────────────────
            # 5. Checklist de cumplimiento
            # ─────────────────────────────────────
            st.markdown("### ✅ Checklist de completitud de acciones")

            tiene_contencion = (
                df_actions[col_acc_cont].notna()
                if col_acc_cont
                else pd.Series(False, index=df_actions.index)
            )
            tiene_corr = (
                df_actions[col_acc_corr].notna()
                if col_acc_corr
                else pd.Series(False, index=df_actions.index)
            )
            tiene_causa = (
                df_actions[col_causa].notna()
                if col_causa
                else pd.Series(False, index=df_actions.index)
            )

            col_ch1, col_ch2, col_ch3 = st.columns(3)
            with col_ch1:
                st.metric("Con acción de contención", int(tiene_contencion.sum()))
            with col_ch2:
                st.metric("Con acción correctiva", int(tiene_corr.sum()))
            with col_ch3:
                st.metric("Con causa raíz definida", int(tiene_causa.sum()))

            # ─────────────────────────────────────
            # 6. Matriz y gráficos de riesgo / prioridad (con porcentajes)
            # ─────────────────────────────────────
            st.markdown("### 📈 Visualización de riesgo y prioridad")

            # 6.1 Matriz cuantitativa Riesgo vs Prioridad
            pivot_riesgo = (
                df_actions.pivot_table(
                    index="Riesgo (automático)",
                    columns="Prioridad (automática)",
                    values="Score riesgo",
                    aggfunc="count",
                    fill_value=0,
                )
                .reindex(index=["Muy alto", "Alto", "Medio", "Bajo"])
            )

            st.markdown("#### Matriz cuantitativa Riesgo vs Prioridad (conteo de CAPA)")
            st.dataframe(pivot_riesgo, use_container_width=True)

            # Distribución por nivel de riesgo con porcentajes
            total_capas = pivot_riesgo.values.sum()
            if total_capas > 0:
                dist_riesgo = (
                    df_actions["Riesgo (automático)"]
                    .value_counts()
                    .reindex(["Muy alto", "Alto", "Medio", "Bajo"])
                    .fillna(0)
                    .astype(int)
                )
                df_dist_riesgo = dist_riesgo.reset_index()
                df_dist_riesgo.columns = ["Riesgo (automático)", "Conteo"]
                df_dist_riesgo["% sobre total"] = (
                    df_dist_riesgo["Conteo"] / total_capas * 100
                ).round(1)

                st.markdown("#### Distribución global por nivel de riesgo")
                st.dataframe(df_dist_riesgo, use_container_width=True)

                st.markdown("**📝 Lectura automática – Matriz Riesgo vs Prioridad**")
                texto_riesgo = []
                for _, r in df_dist_riesgo.iterrows():
                    if r["Conteo"] > 0:
                        texto_riesgo.append(
                            f"- **{r['Riesgo (automático)']}**: {int(r['Conteo'])} CAPA "
                            f"({r['% sobre total']}% del total)."
                        )
                st.write(
                    "La matriz permite ver en qué combinaciones de riesgo y prioridad "
                    "se concentran las incidencias. La distribución actual por nivel de riesgo es:\n"
                    + "\n".join(texto_riesgo)
                )
            else:
                st.caption("No hay datos suficientes para calcular la distribución de riesgo.")

            # 6.2 Gráficos de barras con porcentaje
            try:
                # Distribución de prioridad
                dist_prior = (
                    df_actions["Prioridad (automática)"]
                    .value_counts()
                    .reset_index()
                )
                dist_prior.columns = ["Prioridad (automática)", "Conteo"]
                dist_prior["Porcentaje"] = (
                    dist_prior["Conteo"] / dist_prior["Conteo"].sum() * 100
                ).round(1)
                dist_prior = dist_prior.sort_values("Prioridad (automática)")

                fig_prioridad = px.bar(
                    dist_prior,
                    x="Prioridad (automática)",
                    y="Conteo",
                    text=dist_prior["Porcentaje"].apply(lambda x: f"{x:.1f}%"),
                    title="Distribución de CAPA por prioridad",
                )
                fig_prioridad.update_traces(textposition="outside")
                fig_prioridad.update_layout(uniformtext_minsize=8, uniformtext_mode="hide")

                # Lectura automática para prioridad
                partes_prior = []
                for _, r in dist_prior.iterrows():
                    partes_prior.append(
                        f"**{r['Prioridad (automática)']}**: {int(r['Conteo'])} CAPA "
                        f"({r['Porcentaje']}%)."
                    )
                resumen_prioridad = (
                    "La gráfica muestra cómo se distribuyen las CAPA según la prioridad "
                    "asignada automáticamente:\n" + "\n".join(f"- {p}" for p in partes_prior)
                )

                mostrar_fig_con_lectura(
                    fig_key="plan_prioridad",
                    titulo="Distribución de CAPA por prioridad",
                    fig=fig_prioridad,
                    resumen=resumen_prioridad,
                )

                # Distribución de riesgo
                dist_riesgo_bar = (
                    df_actions["Riesgo (automático)"]
                    .value_counts()
                    .reindex(["Muy alto", "Alto", "Medio", "Bajo"])
                    .fillna(0)
                    .reset_index()
                )
                dist_riesgo_bar.columns = ["Riesgo (automático)", "Conteo"]
                total_riesgo_bar = dist_riesgo_bar["Conteo"].sum()
                if total_riesgo_bar > 0:
                    dist_riesgo_bar["Porcentaje"] = (
                        dist_riesgo_bar["Conteo"] / total_riesgo_bar * 100
                    ).round(1)
                else:
                    dist_riesgo_bar["Porcentaje"] = 0.0

                fig_riesgo = px.bar(
                    dist_riesgo_bar,
                    x="Riesgo (automático)",
                    y="Conteo",
                    text=dist_riesgo_bar["Porcentaje"].apply(lambda x: f"{x:.1f}%"),
                    title="Distribución de CAPA por riesgo",
                )
                fig_riesgo.update_traces(textposition="outside")
                fig_riesgo.update_layout(uniformtext_minsize=8, uniformtext_mode="hide")

                partes_riesgo = []
                for _, r in dist_riesgo_bar.iterrows():
                    partes_riesgo.append(
                        f"**{r['Riesgo (automático)']}**: {int(r['Conteo'])} CAPA "
                        f"({r['Porcentaje']}%)."
                    )
                resumen_riesgo = (
                    "La gráfica muestra la concentración de CAPA según el nivel de riesgo "
                    "estimado automáticamente:\n" + "\n".join(f"- {p}" for p in partes_riesgo)
                )

                mostrar_fig_con_lectura(
                    fig_key="plan_riesgo",
                    titulo="Distribución de CAPA por riesgo",
                    fig=fig_riesgo,
                    resumen=resumen_riesgo,
                )

            except Exception:
                st.caption(
                    "No se pudieron generar los gráficos de Plotly (revisa dependencias o datos)."
                )

            # ─────────────────────────────────────
            # 7. Tabla principal del plan de acciones (vista operativa)
            # ─────────────────────────────────────
            st.markdown("### 📝 Tabla operativa del plan de acciones")

            columnas_base_mostrar = []
            if col_capa:
                columnas_base_mostrar.append(col_capa)
            if col_pedido:
                columnas_base_mostrar.append(col_pedido)
            if col_cliente:
                columnas_base_mostrar.append(col_cliente)
            if col_inc:
                columnas_base_mostrar.append(col_inc)
            if col_causa:
                columnas_base_mostrar.append(col_causa)
            if col_resp:
                columnas_base_mostrar.append(col_resp)
            if col_tecnico:
                columnas_base_mostrar.append(col_tecnico)
            if col_fecha_crea:
                columnas_base_mostrar.append(col_fecha_crea)
            if col_e_repo:
                columnas_base_mostrar.append(col_e_repo)
            if col_e_dev:
                columnas_base_mostrar.append(col_e_dev)
            if col_enlace_8d:
                columnas_base_mostrar.append(col_enlace_8d)

            columnas_inteligentes = [
                "Estado (automático)",
                "Prioridad (automática)",
                "Riesgo (automático)",
                "Score riesgo",
                "Días abiertos (aprox.)",
                "Responsable sugerido",
                "Acción sugerida (texto IA simple)",
            ]

            columnas_finales = columnas_base_mostrar + columnas_inteligentes
            columnas_finales = [c for c in columnas_finales if c in df_actions.columns]

            df_vista = df_actions[columnas_finales].copy()
            st.dataframe(df_vista, use_container_width=True)

            st.markdown("**📝 Lectura automática – Tabla operativa**")
            st.write(
                "La tabla resume, para cada CAPA, la información clave de cliente, pedido, "
                "incidencia, causa raíz, responsables y valores económicos, junto con el "
                "estado y la prioridad calculados automáticamente. La columna de "
                "**acción sugerida** sirve como guía rápida para preparar planes de trabajo "
                "y minutas de reunión, sin necesidad de revisar informe por informe."
            )

            # ─────────────────────────────────────
            # 8. Exportar plan de acciones (Excel)
            # ─────────────────────────────────────
            st.markdown("### 📤 Exportar plan de acciones")

            col_exp1, col_exp2 = st.columns(2)

            with col_exp1:
                buffer_xlsx = BytesIO()
                with pd.ExcelWriter(buffer_xlsx, engine="openpyxl") as writer:
                    df_vista.to_excel(writer, index=False, sheet_name="PlanAcciones")
                buffer_xlsx.seek(0)

                st.download_button(
                    label="⬇️ Descargar plan de acciones (Excel)",
                    data=buffer_xlsx,
                    file_name="Plan_de_Acciones_CAPA.xlsx",
                    mime=(
                        "application/vnd.openxmlformats-officedocument."
                        "spreadsheetml.sheet"
                    ),
                    key="dl_plan_acciones_excel",
                )

# ──────────────────────────────────────────────────────────────
# 10. PIE DE PÁGINA
# ──────────────────────────────────────────────────────────────

st.markdown("---")
st.caption(
    "Sistema de Gestion Operativa de calidad (control y seguimiento de incidencias ECOcero S.L)"
    " Verion 1.0 modelo demostrativo, Neiber Vicney Mendoza" 
    )
